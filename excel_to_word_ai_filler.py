# -*- coding: utf-8 -*-
"""
Excel到Word智能填充工具
使用AI自动匹配列名和字段名
依赖: pip install pandas openpyxl python-docx openai
"""

import pandas as pd
from docx import Document
import os
import json
from openai import OpenAI


# ==================== 配置区域 ====================
OPENAI_API_KEY = ""  # 在这里填入你的OpenAI API Key
OPENAI_BASE_URL = None  # 如果用代理，填入base_url，如 "https://api.openai.com/v1"
MODEL = "gpt-4o-mini"  # 使用的模型，gpt-4o-mini便宜快速
# ================================================


class AIFieldMatcher:
    """AI字段匹配器"""
    
    def __init__(self, api_key=None, base_url=None, model=None):
        """
        初始化
        
        参数:
            api_key: OpenAI API Key
            base_url: API基础URL（代理用）
            model: 使用的模型
        """
        self.api_key = api_key or OPENAI_API_KEY
        self.base_url = base_url or OPENAI_BASE_URL
        self.model = model or MODEL
        
        if not self.api_key:
            raise ValueError("请配置OpenAI API Key")
        
        # 初始化OpenAI客户端
        client_kwargs = {"api_key": self.api_key}
        if self.base_url:
            client_kwargs["base_url"] = self.base_url
        
        self.client = OpenAI(**client_kwargs)
    
    def match_fields(self, excel_columns, word_fields):
        """
        使用AI匹配Excel列名和Word字段名
        
        参数:
            excel_columns: Excel列名列表
            word_fields: Word字段名列表
        
        返回:
            映射字典 {word字段: excel列名}
        """
        prompt = f"""你是一个数据映射专家。请将Excel列名映射到Word字段名。

Excel列名列表：
{json.dumps(excel_columns, ensure_ascii=False, indent=2)}

Word字段名列表：
{json.dumps(word_fields, ensure_ascii=False, indent=2)}

请根据语义相似性进行匹配，返回JSON格式的映射关系：
{{
    "word字段名": "excel列名",
    ...
}}

规则：
1. 根据语义理解进行匹配，比如"零部件名称"可能对应"产品信息"
2. 如果找不到合适的匹配，值设为null
3. 只返回JSON，不要其他解释

返回JSON："""

        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": "你是一个数据映射专家，只返回JSON格式的映射结果。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=1000
        )
        
        result_text = response.choices[0].message.content.strip()
        
        # 提取JSON
        if "```json" in result_text:
            result_text = result_text.split("```json")[1].split("```")[0]
        elif "```" in result_text:
            result_text = result_text.split("```")[1].split("```")[0]
        
        mapping = json.loads(result_text)
        
        # 将null转为None
        mapping = {k: (v if v else None) for k, v in mapping.items()}
        
        return mapping


class ExcelToWordFiller:
    """Excel到Word填充器"""
    
    def __init__(self, excel_path, word_template_path, output_path=None):
        """
        初始化
        
        参数:
            excel_path: Excel文件路径
            word_template_path: Word模板文件路径
            output_path: 输出文件路径，默认在模板同目录
        """
        self.excel_path = excel_path
        self.word_template_path = word_template_path
        self.output_path = output_path or word_template_path.replace('.docx', '_filled.docx')
        
        self.excel_data = None
        self.word_doc = None
        self.field_mapping = {}
    
    def load_excel(self, sheet_name=0):
        """加载Excel文件"""
        # 检测格式
        file_ext = os.path.splitext(self.excel_path)[1].lower()
        engine = 'openpyxl' if file_ext == '.xlsx' else 'xlrd'
        
        try:
            self.excel_data = pd.read_excel(self.excel_path, sheet_name=sheet_name, engine=engine)
        except:
            self.excel_data = pd.read_excel(self.excel_path, sheet_name=sheet_name)
        
        print(f"加载Excel成功: {len(self.excel_data)} 行, {len(self.excel_data.columns)} 列")
        print(f"列名: {list(self.excel_data.columns)}")
        return self.excel_data
    
    def load_word_template(self):
        """加载Word模板"""
        self.word_doc = Document(self.word_template_path)
        print(f"加载Word模板成功")
        return self.word_doc
    
    def extract_word_fields(self):
        """
        提取Word模板中的字段名
        查找 {{字段名}} 格式的占位符
        """
        import re
        fields = set()
        
        # 遍历所有段落
        for para in self.word_doc.paragraphs:
            matches = re.findall(r'\{\{(\w+)\}\}', para.text)
            fields.update(matches)
        
        # 遍历所有表格
        for table in self.word_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{\{(\w+)\}\}', cell.text)
                    fields.update(matches)
        
        print(f"检测到Word字段: {list(fields)}")
        return list(fields)
    
    def ai_match_fields(self, api_key=None, base_url=None, model=None):
        """
        使用AI自动匹配字段
        
        返回:
            映射字典
        """
        excel_columns = list(self.excel_data.columns)
        word_fields = self.extract_word_fields()
        
        print("\n正在使用AI匹配字段...")
        matcher = AIFieldMatcher(api_key, base_url, model)
        self.field_mapping = matcher.match_fields(excel_columns, word_fields)
        
        print("\nAI匹配结果:")
        for word_field, excel_col in self.field_mapping.items():
            status = "✓" if excel_col else "✗"
            print(f"  {status} {word_field} -> {excel_col}")
        
        return self.field_mapping
    
    def manual_adjust_mapping(self):
        """
        手动调整映射（交互式）
        """
        print("\n当前映射关系:")
        for i, (word_field, excel_col) in enumerate(self.field_mapping.items()):
            print(f"  [{i}] {word_field} -> {excel_col or '(未匹配)'}")
        
        print("\n可选操作:")
        print("  1. 输入序号修改映射，如: 0=列名")
        print("  2. 输入 'done' 确认并继续")
        print("  3. 输入 'skip' 跳过调整")
        
        while True:
            user_input = input("\n请输入: ").strip()
            
            if user_input.lower() in ['done', 'skip', '']:
                break
            
            # 解析修改指令
            if '=' in user_input:
                try:
                    idx, new_col = user_input.split('=')
                    idx = int(idx.strip())
                    new_col = new_col.strip()
                    
                    word_field = list(self.field_mapping.keys())[idx]
                    self.field_mapping[word_field] = new_col if new_col else None
                    print(f"  已修改: {word_field} -> {new_col or '(未匹配)'}")
                except:
                    print("  格式错误，请使用: 序号=列名")
        
        return self.field_mapping
    
    def fill_word(self, row_index=0):
        """
        填充Word文档
        
        参数:
            row_index: 使用Excel的第几行数据（默认第一行）
        """
        if not self.field_mapping:
            print("请先执行字段匹配")
            return None
        
        row_data = self.excel_data.iloc[row_index]
        
        # 构建替换字典
        replace_dict = {}
        for word_field, excel_col in self.field_mapping.items():
            if excel_col and excel_col in row_data:
                value = row_data[excel_col]
                replace_dict[f"{{{{{word_field}}}}}"] = str(value) if pd.notna(value) else ""
        
        print(f"\n填充数据: {replace_dict}")
        
        # 替换段落中的占位符
        for para in self.word_doc.paragraphs:
            for key, value in replace_dict.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)
        
        # 替换表格中的占位符
        for table in self.word_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replace_dict.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
        
        # 保存文档
        self.word_doc.save(self.output_path)
        print(f"\nWord文档已保存: {self.output_path}")
        
        return self.output_path


def main():
    """主函数"""
    # ===== 配置区域 =====
    excel_file = r"D:\AI\test_data.xlsx"  # Excel文件路径
    word_template = r"D:\AI\template.docx"  # Word模板路径
    output_file = None  # 输出路径，None表示自动生成
    
    # OpenAI配置
    api_key = ""  # 在这里填入你的API Key
    base_url = None  # 代理地址，如 "https://api.openai.com/v1"
    model = "gpt-4o-mini"
    
    # 是否手动调整映射
    manual_adjust = True  # True=交互式调整，False=全自动
    
    # 使用Excel第几行数据（从0开始）
    data_row = 0
    # ====================
    
    # 检查文件
    if not os.path.exists(excel_file):
        print(f"错误: Excel文件不存在 - {excel_file}")
        return
    
    if not os.path.exists(word_template):
        print(f"错误: Word模板不存在 - {word_template}")
        return
    
    # 创建填充器
    filler = ExcelToWordFiller(excel_file, word_template, output_file)
    
    # 加载文件
    filler.load_excel()
    filler.load_word_template()
    
    # AI匹配字段
    filler.ai_match_fields(api_key, base_url, model)
    
    # 手动调整（可选）
    if manual_adjust:
        filler.manual_adjust_mapping()
    
    # 填充Word
    filler.fill_word(data_row)
    
    print("\n完成！")


if __name__ == '__main__':
    main()
