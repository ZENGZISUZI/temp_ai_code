"""
动态报告生成器
支持从 PDF、Excel、图片提取数据，生成 Word 报告
"""

import os
import json
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field, asdict
from pathlib import Path

# Word 生成
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# 数据提取
import pandas as pd
import pdfplumber  # PDF 文本/表格提取
from PIL import Image  # 图片处理


@dataclass
class ReportConfig:
    """报告配置"""
    title: str = "报告"
    subtitle: str = ""
    author: str = ""
    date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d"))
    company: str = ""
    version: str = "1.0"


@dataclass
class ReportSection:
    """报告章节"""
    title: str
    content: str = ""
    tables: List[List[List[str]]] = field(default_factory=list)
    images: List[str] = field(default_factory=list)
    subsections: List['ReportSection'] = field(default_factory=list)


class DataExtractor:
    """多源数据提取器"""
    
    @staticmethod
    def extract_from_pdf(pdf_path: str) -> Dict[str, Any]:
        """从 PDF 提取文本和表格"""
        result = {
            "text": [],
            "tables": [],
            "images": []
        }
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 提取文本
                text = page.extract_text()
                if text:
                    result["text"].append({
                        "page": page_num + 1,
                        "content": text
                    })
                
                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    result["tables"].append({
                        "page": page_num + 1,
                        "data": table
                    })
        
        return result
    
    @staticmethod
    def extract_from_excel(excel_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """从 Excel 提取数据"""
        result = {
            "sheets": {},
            "all_data": []
        }
        
        # 读取所有 sheet
        xls = pd.ExcelFile(excel_path)
        
        for sheet in xls.sheet_names:
            if sheet_name and sheet != sheet_name:
                continue
            
            df = pd.read_excel(excel_path, sheet_name=sheet)
            
            # 转换为列表格式
            headers = df.columns.tolist()
            data = df.values.tolist()
            
            result["sheets"][sheet] = {
                "headers": headers,
                "data": data
            }
            result["all_data"].append({
                "sheet": sheet,
                "headers": headers,
                "data": data
            })
        
        return result
    
    @staticmethod
    def extract_image_info(image_path: str) -> Dict[str, Any]:
        """获取图片信息"""
        with Image.open(image_path) as img:
            return {
                "path": image_path,
                "width": img.width,
                "height": img.height,
                "format": img.format,
                "mode": img.mode
            }


class WordReportGenerator:
    """Word 报告生成器"""
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        styles = self.doc.styles
        
        # 标题样式
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 0, 0)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # 一级标题
        if 'CustomHeading1' not in [s.name for s in styles]:
            h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(0, 51, 102)
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
        
        # 二级标题
        if 'CustomHeading2' not in [s.name for s in styles]:
            h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(0, 76, 153)
            h2_style.paragraph_format.space_before = Pt(10)
            h2_style.paragraph_format.space_after = Pt(4)
    
    def add_cover(self):
        """添加封面"""
        # 空行
        for _ in range(3):
            self.doc.add_paragraph()
        
        # 标题
        title = self.doc.add_paragraph(self.config.title, style='CustomTitle')
        
        # 副标题
        if self.config.subtitle:
            subtitle = self.doc.add_paragraph(self.config.subtitle)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(14)
        
        # 空行
        for _ in range(5):
            self.doc.add_paragraph()
        
        # 元信息
        meta_info = [
            f"作者：{self.config.author}" if self.config.author else None,
            f"日期：{self.config.date}",
            f"版本：{self.config.version}" if self.config.version else None,
            f"单位：{self.config.company}" if self.config.company else None
        ]
        
        for info in meta_info:
            if info:
                p = self.doc.add_paragraph(info)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 分页
        self.doc.add_page_break()
    
    def add_toc(self):
        """添加目录（占位符，需 Word 手动更新）"""
        p = self.doc.add_paragraph("目录")
        p.style = 'CustomHeading1'
        
        p = self.doc.add_paragraph("（在 Word 中右键更新目录）")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        self.doc.add_page_break()
    
    def add_section(self, section: ReportSection, level: int = 1):
        """添加章节"""
        # 标题
        style_name = f'CustomHeading{min(level, 2)}'
        self.doc.add_paragraph(section.title, style=style_name)
        
        # 正文内容
        if section.content:
            for para in section.content.split('\n'):
                if para.strip():
                    p = self.doc.add_paragraph(para.strip())
        
        # 表格
        for table_data in section.tables:
            self._add_table(table_data)
        
        # 图片
        for img_path in section.images:
            if os.path.exists(img_path):
                self._add_image(img_path)
        
        # 子章节
        for subsection in section.subsections:
            self.add_section(subsection, level + 1)
    
    def _add_table(self, data: List[List[str]]):
        """添加表格"""
        if not data:
            return
        
        rows = len(data)
        cols = len(data[0])
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data) if cell_data else ""
                
                # 表头加粗
                if i == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # 表格后空行
        self.doc.add_paragraph()
    
    def _add_image(self, image_path: str, width: float = 5.0):
        """添加图片"""
        try:
            p = self.doc.add_paragraph()
            run = p.add_run()
            run.add_picture(image_path, width=Inches(width))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            self.doc.add_paragraph(f"[图片加载失败: {image_path}]")
    
    def add_paragraph(self, text: str, style: str = None):
        """添加段落"""
        p = self.doc.add_paragraph(text, style=style)
        return p
    
    def add_bullet_list(self, items: List[str]):
        """添加无序列表"""
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')
    
    def add_numbered_list(self, items: List[str]):
        """添加有序列表"""
        for item in items:
            self.doc.add_paragraph(item, style='List Number')
    
    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)
        print(f"报告已生成: {output_path}")


class ReportBuilder:
    """报告构建器 - 整合多源数据"""
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.generator = WordReportGenerator(config)
        self.sections: List[ReportSection] = []
        self.extractor = DataExtractor()
    
    def load_from_pdf(self, pdf_path: str, section_title: str = "PDF 数据") -> 'ReportBuilder':
        """从 PDF 加载数据并创建章节"""
        data = self.extractor.extract_from_pdf(pdf_path)
        
        section = ReportSection(title=section_title)
        
        # 添加文本内容
        for page_data in data["text"]:
            section.content += f"\n[第 {page_data['page']} 页]\n{page_data['content']}\n"
        
        # 添加表格
        for table_data in data["tables"]:
            section.tables.append(table_data["data"])
        
        self.sections.append(section)
        return self
    
    def load_from_excel(self, excel_path: str, section_title: str = "Excel 数据") -> 'ReportBuilder':
        """从 Excel 加载数据并创建章节"""
        data = self.extractor.extract_from_excel(excel_path)
        
        section = ReportSection(title=section_title)
        
        for sheet_data in data["all_data"]:
            # 创建子章节
            subsection = ReportSection(
                title=f"工作表: {sheet_data['sheet']}",
                tables=[[sheet_data['headers']] + list(sheet_data['data'])]
            )
            section.subsections.append(subsection)
        
        self.sections.append(section)
        return self
    
    def add_images(self, image_paths: List[str], section_title: str = "图片") -> 'ReportBuilder':
        """添加图片章节"""
        section = ReportSection(
            title=section_title,
            images=image_paths
        )
        self.sections.append(section)
        return self
    
    def add_custom_section(self, section: ReportSection) -> 'ReportBuilder':
        """添加自定义章节"""
        self.sections.append(section)
        return self
    
    def build(self, output_path: str, with_cover: bool = True, with_toc: bool = True):
        """构建报告"""
        # 封面
        if with_cover:
            self.generator.add_cover()
        
        # 目录
        if with_toc:
            self.generator.add_toc()
        
        # 各章节
        for section in self.sections:
            self.generator.add_section(section)
        
        # 保存
        self.generator.save(output_path)


# ==================== 使用示例 ====================

def example_usage():
    """使用示例"""
    
    # 1. 配置报告
    config = ReportConfig(
        title="项目测试报告",
        subtitle="自动化测试结果汇总",
        author="测试团队",
        company="XX科技有限公司",
        version="1.0"
    )
    
    # 2. 创建构建器
    builder = ReportBuilder(config)
    
    # 3. 从各数据源加载数据
    # builder.load_from_pdf("input/test_report.pdf", "PDF 测试数据")
    # builder.load_from_excel("input/test_data.xlsx", "Excel 测试数据")
    # builder.add_images(["input/screenshot1.png", "input/screenshot2.png"], "测试截图")
    
    # 4. 添加自定义章节
    summary_section = ReportSection(
        title="测试概述",
        content="""
        本次测试覆盖了系统的核心功能模块，包括：
        - 用户登录/注册功能
        - 数据查询功能
        - 报表生成功能
        
        测试环境：Windows 10 / Chrome 120
        测试时间：2024-01-15 至 2024-01-20
        """,
        subsections=[
            ReportSection(
                title="测试结果统计",
                tables=[
                    ["测试类型", "用例数", "通过", "失败", "通过率"],
                    ["功能测试", "100", "95", "5", "95%"],
                    ["性能测试", "50", "48", "2", "96%"],
                    ["安全测试", "30", "28", "2", "93.3%"]
                ]
            )
        ]
    )
    builder.add_custom_section(summary_section)
    
    # 5. 生成报告
    builder.build("output/report.docx")
    print("报告生成完成！")


if __name__ == "__main__":
    example_usage()
