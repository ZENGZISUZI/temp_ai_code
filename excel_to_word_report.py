# -*- coding: utf-8 -*-
"""
Excel测试报告转Word文档工具
自动提取Excel数据生成标准Word报告模板
依赖: pip install pandas openpyxl python-docx xlrd
      - openpyxl: 支持 .xlsx 格式
      - xlrd: 支持 .xls 格式（旧版Excel）
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os
import re
import traceback
from datetime import datetime
from openpyxl import load_workbook
import xlrd

# 尝试导入win32com（可选依赖）
try:
    import win32com.client
except ImportError:
    win32com = None


# ==================== 智能字段映射配置 ====================

# 概述部分字段映射（Excel列名 -> Word报告字段）
OVERVIEW_FIELD_MAPPING = {
    # 产品信息相关
    '产品信息': ['零部件名称', '产品名称', '产品型号', '设备名称', '名称'],
    # 试验信息相关
    '试验信息': ['试验名称', '试验类型', '测试类型', '试验目的'],
    # 工作模式相关
    '工作模式': ['工作模式', '运行模式', '工作状态', '模式'],
    # 测试仪器设备相关
    '测试仪器设备': ['测试仪器', '仪器设备', '试验设备', '设备', '使用仪器'],
    # 封面专用字段
    '产品名称': ['产品名称', '零部件名称', '设备名称', '名称'],
    '零件号': ['零件号', '零件编号', '部件号', '图号'],
    '客户名称': ['客户名称', '客户', '委托单位', '委托方'],
    '客户地址': ['客户地址', '地址', '委托单位地址'],
    '收样日期': ['收样日期', '收样时间', '接收日期', '样品接收日期'],
    '试验日期': ['试验日期', '测试日期', '试验时间', '测试时间'],
    '测试结论': ['测试结论', '试验结论', '结论', '测试结果'],
    '产品型号': ['产品型号', '型号', '规格型号', '设备型号'],
}

# 小用例表格字段映射（Word字段 -> Excel可能列名）
TESTCASE_FIELD_MAPPING = {
    '开始日期': ['开始日期', '开始时间', '起始日期', '开始'],
    '结束日期': ['结束日期', '结束时间', '终止日期', '结束', '计划完成时间', '计划完成日期'],
    '样机数量': ['样机数量', '数量', '样品数量', '台数', 'DV数量'],
    '样机编号': ['样机编号', '编号', '样品编号', '机号'],
    '试验机构': ['试验机构', '检测机构', '测试机构', '机构'],
    '试验环境': ['试验环境', '环境条件', '环境', '测试环境'],
    '试验标准': ['试验标准', '标准', '测试标准', '参考标准'],
    '试验条件': ['试验条件', '条件', '测试条件', '试验方法'],
    '规格要求': ['规格要求', '要求', '技术要求', '规格'],
    '试验数据': ['试验数据', '数据', '测试数据', '结果数据'],
    '试验结论': ['试验结论', '结论', '测试结论', '结果'],
}


def clean_case_number(name):
    """
    清理用例名字中的数字前缀
    
    支持格式:
    - "1、xxx" -> "xxx"
    - "一、xxx" -> "xxx"
    - "1.xxx" -> "xxx"
    - "1 xxx" -> "xxx"
    - "（1）xxx" -> "xxx"
    - "(1)xxx" -> "xxx"
    
    参数:
        name: 用例名字
        
    返回:
        清理后的名字
    """
    if not name:
        return name
    
    # 中文数字映射
    chinese_nums = '一二三四五六七八九十'
    
    # 匹配模式：数字/中文数字 + 标点符号
    patterns = [
        r'^[\d]+\s*[、.．。:：]\s*',  # 1、 1. 1．
        r'^[一二三四五六七八九十]+\s*[、.．。:：]\s*',  # 一、 一.
        r'^[\(（][\d]+[）\)]\s*',  # (1) （1）
        r'^[\d]+\s+',  # 1 开头加空格
    ]
    
    result = name.strip()
    for pattern in patterns:
        result = re.sub(pattern, '', result)
    
    return result.strip()


def extract_case_number(name):
    """
    提取用例名字中的数字前缀
    
    参数:
        name: 用例名字
        
    返回:
        数字（阿拉伯数字），如果没有则返回None
    """
    if not name:
        return None
    
    # 中文数字映射
    chinese_to_num = {
        '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
        '六': 6, '七': 7, '八': 8, '九': 9, '十': 10
    }
    
    # 尝试匹配中文数字
    match = re.match(r'^([一二三四五六七八九十]+)\s*[、.．。:：]', name)
    if match:
        chinese_num = match.group(1)
        return chinese_to_num.get(chinese_num)
    
    # 尝试匹配阿拉伯数字
    match = re.match(r'^[\(（]?(\d+)[）\)]?\s*[、.．。:：]?', name)
    if match:
        return int(match.group(1))
    
    return None


def validate_chapter_position(big_cases):
    """
    验证大用例是否放在正确的章节位置
    
    参数:
        big_cases: 大用例列表
        
    返回:
        验证结果列表 [{'name': 'xxx', 'expected': 1, 'actual': 1, 'valid': True}, ...]
    """
    results = []
    
    for idx, big_case in enumerate(big_cases, 1):
        name = big_case.get('name', '')
        expected_num = extract_case_number(name)
        
        result = {
            'name': name,
            'expected': expected_num,
            'actual': idx,
            'valid': expected_num is None or expected_num == idx
        }
        
        if not result['valid']:
            print(f"⚠️ 警告: 大用例 '{name}' 位置不匹配，期望在第{expected_num}章节，实际在第{idx}章节")
        
        results.append(result)
    
    return results


def find_best_match(target_field, excel_columns, mapping_dict):
    """
    智能匹配：根据目标字段在Excel列名中找最佳匹配

    参数:
        target_field: 目标字段名（如"产品信息"）
        excel_columns: Excel实际列名列表
        mapping_dict: 字段映射字典

    返回:
        匹配到的Excel列名，未匹配返回None
    """
    if target_field not in mapping_dict:
        return None

    keywords = mapping_dict[target_field]

    for keyword in keywords:
        for col in excel_columns:
            if col and keyword in str(col):
                return col

    return None


def smart_match_field(field_name, excel_columns):
    """
    智能匹配单个字段

    参数:
        field_name: 字段名
        excel_columns: Excel列名列表

    返回:
        匹配到的列名
    """
    # 先尝试精确匹配
    if field_name in excel_columns:
        return field_name

    # 再尝试包含匹配
    for col in excel_columns:
        if col and field_name in str(col):
            return col

    # 尝试反向匹配（Excel列名包含在字段名中）
    for col in excel_columns:
        if col and str(col) in field_name:
            return col

    return None


def set_cell_font(cell, font_name='微软雅黑', font_size=10.5, bold=False, align_center=True, vertical_center=True):
    """
    设置单元格字体和对齐方式
    
    参数:
        cell: 单元格对象
        font_name: 字体名称（默认微软雅黑）
        font_size: 字体大小（默认五号10.5pt）
        bold: 是否加粗
        align_center: 是否水平居中对齐（False则靠左）
        vertical_center: 是否垂直居中
    """
    # 设置字体
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        # 设置段落水平对齐
        if align_center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 设置单元格垂直居中
    if vertical_center:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_border(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_cover_page(doc, report_name, report_number, company_name="公司", company_full_name=None, company_address=None, logo_path=None, font_config=None, overview_data=None):
    """
    添加报告封面页
    
    参数:
        doc: Word文档对象
        report_name: 报告名称
        report_number: 报告编号
        company_name: 公司名称简称
        company_full_name: 公司名称全称（Logo下方显示）
        company_address: 公司地址（公司全称下方显示）
        logo_path: Logo图片路径
        font_config: 字体配置
        overview_data: 概述数据（用于封面表格）
    """
    fc = font_config or {}
    font_name = fc.get('font_name', '微软雅黑')
    body_size = fc.get('body_size', 10.5)
    
    # ===== 顶部信息表格（4行4列）=====
    # 属性排列：产品名称|值|零件号|值 / 客户名称|值|客户地址|值 / 收样日期|值|试验日期|值 / 测试结论|值|产品型号|值
    cover_fields = [
        ('产品名称', '零件号'),
        ('客户名称', '客户地址'),
        ('收样日期', '试验日期'),
        ('测试结论', '产品型号'),
    ]
    
    cover_table = doc.add_table(rows=4, cols=4)
    cover_table.alignment = WD_TABLE_ALIGNMENT.LEFT  # 靠左对齐
    set_table_border(cover_table)
    
    # 设置列宽
    col_widths = [Cm(3), Cm(5), Cm(3), Cm(5)]  # 属性列窄，值列宽
    for i, width in enumerate(col_widths):
        cover_table.columns[i].width = Cm(width)
    
    # 填充表格
    data = overview_data or {}
    for row_idx, (field1, field2) in enumerate(cover_fields):
        # 第一列：属性1
        cell = cover_table.cell(row_idx, 0)
        cell.text = field1
        set_cell_font(cell, font_name=font_name, font_size=body_size, bold=True, align_center=False, vertical_center=True)
        
        # 第二列：值1
        cell = cover_table.cell(row_idx, 1)
        value1 = data.get(field1, '')
        cell.text = str(value1) if value1 else ''
        set_cell_font(cell, font_name=font_name, font_size=body_size, bold=False, align_center=False, vertical_center=True)
        
        # 第三列：属性2
        cell = cover_table.cell(row_idx, 2)
        cell.text = field2
        set_cell_font(cell, font_name=font_name, font_size=body_size, bold=True, align_center=False, vertical_center=True)
        
        # 第四列：值2
        cell = cover_table.cell(row_idx, 3)
        value2 = data.get(field2, '')
        cell.text = str(value2) if value2 else ''
        set_cell_font(cell, font_name=font_name, font_size=body_size, bold=False, align_center=False, vertical_center=True)
    
    # 添加空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 报告名称（作为封面标题，从配置获取）
    if report_name:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(report_name)
        title_run.font.name = font_name
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        title_run.font.size = Pt(36)
        title_run.font.bold = True
    
    # ===== 签名表格（3行4列）=====
    # 第3列和第4列合并三行
    # 行1: 编制 | 值 | 签发日期 | 值
    # 行2: 审核 | 值 | (合并)  | (合并)
    # 行3: 批准 | 值 | (合并)  | (合并)
    sign_table = doc.add_table(rows=3, cols=4)
    sign_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(sign_table)
    
    # 设置列宽
    sign_col_widths = [Cm(2.5), Cm(4), Cm(2.5), Cm(4)]
    for i, width in enumerate(sign_col_widths):
        sign_table.columns[i].width = Cm(width)
    
    # 合并第3列（签发日期）的三行
    sign_table.cell(0, 2).merge(sign_table.cell(1, 2)).merge(sign_table.cell(2, 2))
    # 合并第4列（值）的三行
    sign_table.cell(0, 3).merge(sign_table.cell(1, 3)).merge(sign_table.cell(2, 3))
    
    # 行1: 编制 | 值 | 签发日期 | 值
    cell = sign_table.cell(0, 0)
    cell.text = '编制'
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=True, align_center=False, vertical_center=True)
    
    cell = sign_table.cell(0, 1)
    cell.text = ''
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=False, align_center=False, vertical_center=True)
    
    cell = sign_table.cell(0, 2)
    cell.text = '签发日期'
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=True, align_center=False, vertical_center=True)
    
    cell = sign_table.cell(0, 3)
    cell.text = ''
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=False, align_center=False, vertical_center=True)
    
    # 行2: 审核 | 值 | (合并) | (合并)
    cell = sign_table.cell(1, 0)
    cell.text = '审核'
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=True, align_center=False, vertical_center=True)
    
    cell = sign_table.cell(1, 1)
    cell.text = ''
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=False, align_center=False, vertical_center=True)
    
    # 行3: 批准 | 值 | (合并) | (合并)
    cell = sign_table.cell(2, 0)
    cell.text = '批准'
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=True, align_center=False, vertical_center=True)
    
    cell = sign_table.cell(2, 1)
    cell.text = ''
    set_cell_font(cell, font_name=font_name, font_size=body_size, bold=False, align_center=False, vertical_center=True)
    
    # Logo图片（居中）
    if logo_path and os.path.exists(logo_path):
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Cm(4), height=Cm(2.6))  # 较大的logo
        except Exception as e:
            print(f"警告: 无法添加Logo图片 - {e}")
    
    # 公司名称全称（Logo下方）
    if company_full_name:
        company_full_para = doc.add_paragraph()
        company_full_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_full_run = company_full_para.add_run(company_full_name)
        company_full_run.font.name = font_name
        company_full_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        company_full_run.font.size = Pt(14)
    
    # 公司地址（公司全称下方）
    if company_address:
        addr_para = doc.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr_run = addr_para.add_run(f'公司地址：{company_address}')
        addr_run.font.name = font_name
        addr_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        addr_run.font.size = Pt(12)
    
    # 版权声明
    copyright_para = doc.add_paragraph()
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_run = copyright_para.add_run('版权所有 侵权必究')
    copyright_run.font.name = font_name
    copyright_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    copyright_run.font.size = Pt(10)
    
    # 分页符
    doc.add_page_break()


def add_declaration_page(doc, company_name="公司", font_config=None):
    """
    添加声明页
    
    参数:
        doc: Word文档对象
        company_name: 公司名称
        font_config: 字体配置
    """
    fc = font_config or {}
    font_name = fc.get('font_name', '微软雅黑')
    
    # 声明标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('声  明')
    title_run.font.name = font_name
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    
    # 空行
    doc.add_paragraph()
    
    # 声明内容
    declaration_text = f"""1. 本报告无检测专用章、骑缝章无效。

2. 本报告无主检、审核、批准签字无效。

3. 本报告涂改、复印、扫描无效。

4. 本报告仅对送检样品负责。

5. 未经{company_name}书面批准，不得部分复制本报告。

6. 如对本报告有异议，请在收到报告之日起15日内向{company_name}提出。"""
    
    content_para = doc.add_paragraph()
    content_run = content_para.add_run(declaration_text)
    content_run.font.name = font_name
    content_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    content_run.font.size = Pt(12)
    content_para.paragraph_format.line_spacing = 1.5
    
    # 分页符
    doc.add_page_break()


def add_header_footer(doc, report_name, report_number, logo_path=None, company_name="公司"):
    """
    添加页眉和页脚（带横线分隔）
    
    参数:
        doc: Word文档对象
        report_name: 报告名称（文件名）
        report_number: 报告编号
        logo_path: Logo图片路径（可选）
        company_name: 公司名称（用于保密信息）
    """
    # 获取文档的第一个节
    section = doc.sections[0]
    
    # ===== 设置页眉 =====
    header = section.header
    header.is_linked_to_previous = False
    
    # 删除默认段落（避免页眉上方出现回车符号）
    for para in header.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    # 创建页眉表格（2列：Logo+报告名 | 报告编号）
    header_table = header.add_table(rows=1, cols=2, width=Inches(7.5))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置列宽（各占一半）
    header_table.columns[0].width = Inches(3.75)  # Logo + 报告名
    header_table.columns[1].width = Inches(3.75)  # 报告编号
    
    # 第1列：Logo + 报告名称（左对齐）
    cell_left = header_table.cell(0, 0)
    para_left = cell_left.paragraphs[0]
    para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_left.paragraph_format.space_before = Pt(0)
    para_left.paragraph_format.space_after = Pt(0)
    para_left.paragraph_format.line_spacing = 1.0
    
    # 添加Logo图片
    if logo_path and os.path.exists(logo_path):
        try:
            run_logo = para_left.add_run()
            run_logo.add_picture(logo_path, width=Cm(0.75), height=Cm(0.49))
            para_left.add_run("  ")
        except Exception as e:
            print(f"警告: 无法添加Logo图片 - {e}")
    
    # 添加报告名称
    run_name = para_left.add_run(report_name)
    run_name.font.name = '微软雅黑'
    run_name.font.size = Pt(9)
    run_name._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 第2列：报告编号（右对齐）
    cell_right = header_table.cell(0, 1)
    para_right = cell_right.paragraphs[0]
    para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_right.paragraph_format.space_before = Pt(0)
    para_right.paragraph_format.space_after = Pt(0)
    para_right.paragraph_format.line_spacing = 1.0
    run_number = para_right.add_run(f"报告编号:{report_number}")
    run_number.font.name = '微软雅黑'
    run_number.font.size = Pt(9)
    run_number._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 设置表格只有底部边框（横线紧贴文字）
    set_table_border_with_bottom_line(header_table)
    
    # ===== 设置页脚 =====
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # 删除默认段落（避免页脚上方出现回车符号）
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    # 创建页脚表格（2列：保密信息 | 页码）
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(7.5))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置列宽（保密信息列宽，页码列窄）
    footer_table.columns[0].width = Inches(5.0)  # 保密信息
    footer_table.columns[1].width = Inches(2.5)  # 页码
    
    # 第1列：保密信息（居中）
    cell_secret = footer_table.cell(0, 0)
    para_secret = cell_secret.paragraphs[0]
    para_secret.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_secret.paragraph_format.space_before = Pt(0)
    para_secret.paragraph_format.space_after = Pt(0)
    para_secret.paragraph_format.line_spacing = 1.0
    run_secret = para_secret.add_run(f"{company_name}保密信息，未经授权禁止扩散！")
    run_secret.font.name = '微软雅黑'
    run_secret.font.size = Pt(9)
    run_secret.font.bold = True
    run_secret._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 第2列：页码（右对齐）
    cell_page = footer_table.cell(0, 1)
    para_page = cell_page.paragraphs[0]
    para_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_page.paragraph_format.space_before = Pt(0)
    para_page.paragraph_format.space_after = Pt(0)
    para_page.paragraph_format.line_spacing = 1.0
    
    # 添加页码字段
    run_page = para_page.add_run("第 ")
    run_page.font.name = '微软雅黑'
    run_page.font.size = Pt(9)
    run_page._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_page_number_field(para_page, '微软雅黑', 9)
    
    run_page2 = para_page.add_run(" 页，共 ")
    run_page2.font.name = '微软雅黑'
    run_page2.font.size = Pt(9)
    run_page2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    add_total_pages_field(para_page, '微软雅黑', 9)
    
    run_page3 = para_page.add_run(" 页")
    run_page3.font.name = '微软雅黑'
    run_page3.font.size = Pt(9)
    run_page3._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 设置表格只有顶部边框（横线紧贴文字）
    set_table_border_with_top_line(footer_table)
    
    # 删除表格后面可能自动添加的空段落
    for para in footer.paragraphs:
        if not para.text.strip():
            p = para._element
            p.getparent().remove(p)


def set_table_border_with_bottom_line(table):
    """
    设置表格只有底部边框线（用于页眉）
    
    参数:
        table: 表格对象
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    # 只设置底部边框
    for border_name in ['top', 'left', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    # 底部边框显示
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_table_border_with_top_line(table):
    """
    设置表格只有顶部边框线（用于页脚）
    
    参数:
        table: 表格对象
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    # 只设置顶部边框
    for border_name in ['bottom', 'left', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    # 顶部边框显示
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_table_border(table, show_border=True):
    """
    设置表格边框
    
    参数:
        table: 表格对象
        show_border: 是否显示边框，False则隐藏边框
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        if show_border:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
        else:
            border.set(qn('w:val'), 'nil')
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_page_number_field(paragraph, font_name='微软雅黑', font_size=9):
    """添加当前页码字段"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # 设置字体
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_total_pages_field(paragraph, font_name='微软雅黑', font_size=9):
    """添加总页数字段"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # 设置字体
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_watermark_to_docx(doc, watermark_text):
    """
    为文档添加斜向水印（从左下角到右上角）
    水印在底层，正文内容浮在上面
    
    参数:
        doc: Word文档对象
        watermark_text: 水印文字
    """
    # VML和Office命名空间
    VML_NS = 'urn:schemas-microsoft-com:vml'
    OFFICE_NS = 'urn:schemas-microsoft-com:office:office'
    
    # 为每个节添加水印
    for section in doc.sections:
        header = section.header
        
        # 获取或创建header的XML元素
        header_elem = header._element
        
        # 创建水印段落
        watermark_para = OxmlElement('w:p')
        
        # 创建段落属性
        pPr = OxmlElement('w:pPr')
        watermark_para.append(pPr)
        
        # 创建VML形状（斜向水印）
        # 使用带命名空间的XML字符串
        shape_xml = f'''
        <v:shape xmlns:v="urn:schemas-microsoft-com:vml" 
                 xmlns:o="urn:schemas-microsoft-com:office:office"
                 id="Watermark" 
                 style="position:absolute;margin-left:0;margin-top:0;width:400pt;height:80pt;rotation:315;z-index:-251657216;mso-position-horizontal:center;mso-position-vertical:center;mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
                 coordsize="21600,21600" 
                 allowincell="f" 
                 filled="t" 
                 stroked="f">
            <v:fill opacity="0.3" on="t"/>
            <v:textpath style="font-family:&quot;Arial&quot;;font-size:36pt" on="t" string="{watermark_text}"/>
        </v:shape>
        '''
        
        # 解析XML并添加到段落
        shape_elem = etree.fromstring(shape_xml)
        watermark_para.append(shape_elem)
        
        # 将水印段落添加到header开头
        header_elem.insert(0, watermark_para)


def add_heading_with_number(doc, text, level=1, font_config=None):
    """
    添加带编号的标题（使用Word标题样式，支持目录和导航窗格）
    
    参数:
        doc: Word文档对象
        text: 标题文字
        level: 标题级别（1=大标题，2=次标题，3=小标题）
        font_config: 字体配置字典
    """
    # 默认配置
    default_config = {
        'font_name': '微软雅黑',
        'title1_size': 16,
        'title1_bold': True,
        'title2_size': 12,
        'title2_bold': True,
        'title3_size': 12,
        'title3_bold': False,
    }
    config = default_config.copy()
    if font_config:
        config.update(font_config)
    
    # 使用Word标题样式（支持目录和导航窗格）
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 设置字体格式
    for run in heading.runs:
        run.font.name = config['font_name']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_name'])
        
        if level == 1:
            run.font.size = Pt(config['title1_size'])
            run.font.bold = config['title1_bold']
        elif level == 2:
            run.font.size = Pt(config['title2_size'])
            run.font.bold = config['title2_bold']
        else:
            run.font.size = Pt(config['title3_size'])
            run.font.bold = config['title3_bold']
    
    return heading


def close_word_document(file_path):
    """
    关闭占用指定文件的Word文档
    
    参数:
        file_path: 文件路径
        
    返回:
        True: 成功关闭或文件未被占用
        False: 关闭失败
    """
    if win32com is None:
        return True
    
    try:
        # 获取已运行的Word实例
        word = win32com.client.Dispatch("Word.Application")
        
        abs_path = os.path.abspath(file_path)
        
        # 遍历所有打开的文档
        for doc in word.Documents:
            try:
                # 比较文件路径
                doc_path = os.path.abspath(doc.FullName)
                if doc_path.lower() == abs_path.lower():
                    print(f"  检测到文件被占用，正在关闭: {os.path.basename(file_path)}")
                    doc.Close(SaveChanges=False)  # 不保存更改
                    print(f"  ✓ 已关闭占用文件")
                    break
            except:
                continue
                
        return True
    except Exception as e:
        # Word未运行或其他错误，文件未被占用
        return True


def update_toc_in_word(word_path):
    """
    使用win32com打开Word文档并更新目录
    
    参数:
        word_path: Word文档路径
    """
    if win32com is None:
        print("⚠️ 未安装win32com，无法自动更新目录")
        print("  提示: pip install pywin32")
        return False
    
    try:
        # 使用DispatchEx创建独立的Word实例，避免影响用户已打开的Word
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False  # 后台运行
        
        try:
            # 打开文档
            doc = word.Documents.Open(os.path.abspath(word_path))
            
            # 方法1：全选后更新所有域
            word.Selection.WholeStory()  # 全选
            word.Selection.Fields.Update()  # 更新域
            
            # 方法2：遍历所有故事范围更新域
            for story in doc.StoryRanges:
                story.Fields.Update()
            
            # 保存并关闭文档
            doc.Save()
            doc.Close()
            
            print("✓ 目录已自动更新")
            return True
        finally:
            # 退出独立的Word实例
            word.Quit()
            
    except Exception as e:
        print(f"⚠️ 更新目录失败: {e}")
        return False


def add_bookmark(paragraph, bookmark_name):
    """
    为段落添加书签
    
    参数:
        paragraph: 段落对象
        bookmark_name: 书签名称
    """
    tag = paragraph._p
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    
    tag.insert(0, bookmark_start)
    tag.append(bookmark_end)


def add_hyperlink(paragraph, text, bookmark_name, font_name='微软雅黑', font_size=10.5, bold=False):
    """
    添加超链接到段落（指向书签）
    
    参数:
        paragraph: 段落对象
        text: 显示文字
        bookmark_name: 目标书签名称
        font_name: 字体名称
        font_size: 字体大小
        bold: 是否加粗
    """
    # 创建超链接元素
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    # 创建run元素
    new_run = OxmlElement('w:r')
    
    # 设置字体
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    
    # 设置字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))  # Word字号单位是半磅
    rPr.append(sz)
    
    # 设置加粗
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    
    new_run.append(rPr)
    
    # 设置文字
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def setup_heading_styles(doc, font_config=None):
    """
    设置Word标题样式（确保导航窗格能识别）
    
    参数:
        doc: Word文档对象
        font_config: 字体配置字典
    """
    
    default_config = {
        'font_name': '微软雅黑',
        'title1_size': 16,
        'title1_bold': True,
        'title2_size': 12,
        'title2_bold': True,
        'title3_size': 12,
        'title3_bold': False,
    }
    config = default_config.copy()
    if font_config:
        config.update(font_config)
    
    # 设置 Heading 1 样式
    style1 = doc.styles['Heading 1']
    style1.font.name = config['font_name']
    style1.font.size = Pt(config['title1_size'])
    style1.font.bold = config['title1_bold']
    style1.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    style1._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_name'])
    
    # 设置 Heading 2 样式
    style2 = doc.styles['Heading 2']
    style2.font.name = config['font_name']
    style2.font.size = Pt(config['title2_size'])
    style2.font.bold = config['title2_bold']
    style2.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    style2._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_name'])
    
    # 设置 Heading 3 样式
    style3 = doc.styles['Heading 3']
    style3.font.name = config['font_name']
    style3.font.size = Pt(config['title3_size'])
    style3.font.bold = config['title3_bold']
    style3.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    style3._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_name'])


def add_body_paragraph(doc, text, font_config=None):
    """
    添加正文段落
    
    参数:
        doc: Word文档对象
        text: 正文内容
        font_config: 字体配置字典
    """
    default_config = {
        'font_name': '微软雅黑',
        'body_size': 10.5,
        'body_bold': False,
    }
    config = default_config.copy()
    if font_config:
        config.update(font_config)
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = para.add_run(text)
    run.font.name = config['font_name']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_name'])
    run.font.size = Pt(config['body_size'])
    run.font.bold = config['body_bold']
    
    return para


def format_date_only(value):
    """
    格式化日期，只保留日期部分（去掉时分秒）
    
    参数:
        value: 日期值（可能是字符串、datetime对象等）
        
    返回:
        只包含日期的字符串
    """
    if not value:
        return ''
    
    value_str = str(value).strip()
    
    # 匹配 YYYY-MM-DD 或 YYYY/MM/DD 格式
    date_match = re.match(r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})', value_str)
    if date_match:
        return date_match.group(1).replace('/', '-')
    
    # 匹配 YYYYMMDD 格式
    date_match = re.match(r'(\d{4})(\d{2})(\d{2})', value_str)
    if date_match:
        return f"{date_match.group(1)}-{date_match.group(2)}-{date_match.group(3)}"
    
    # 如果是 datetime 对象
    try:
        if isinstance(value, datetime):
            return value.strftime('%Y-%m-%d')
    except:
        pass
    
    # 返回原始值
    return value_str


def format_quantity(value):
    """
    格式化数量，如果只有数字则添加pcs单位
    
    参数:
        value: 数量值
        
    返回:
        格式化后的数量字符串
    """
    if not value:
        return ''
    
    value_str = str(value).strip()
    
    # 如果已经是纯数字，添加pcs
    if re.match(r'^\d+(\.\d+)?$', value_str):
        return f"{value_str} pcs"
    
    # 如果已经包含单位（如 pcs、台、个等），直接返回
    return value_str


def format_sample_number(value):
    """
    格式化样机编号，将 1/2/3 转换为 1#、2#、3#
    
    参数:
        value: 样机编号值
        
    返回:
        格式化后的样机编号字符串
    """
    if not value:
        return ''
    
    value_str = str(value).strip()
    
    # 处理 1/2/3 格式（用斜杠分隔）
    if '/' in value_str:
        parts = value_str.split('/')
        formatted_parts = [f"{p.strip()}#" for p in parts if p.strip()]
        return '、'.join(formatted_parts)
    
    # 处理 1,2,3 格式（用逗号分隔）
    if ',' in value_str:
        parts = value_str.split(',')
        formatted_parts = [f"{p.strip()}#" for p in parts if p.strip()]
        return '、'.join(formatted_parts)
    
    # 处理单个数字
    if re.match(r'^\d+$', value_str):
        return f"{value_str}#"
    
    # 其他格式直接返回
    return value_str


def create_testcase_table(doc, data_dict, font_config=None, col_widths=None):
    """
    创建测试用例表格
    
    参数:
        doc: Word文档对象
        data_dict: 数据字典 {字段名: 值}
        font_config: 字体配置字典
        col_widths: 列宽列表（单位：厘米），如 [2.5, 4, 2.5, 4]
    """
    # 默认配置
    default_config = {
        'font_name': '微软雅黑',
        'body_size': 10.5,
    }
    config = default_config.copy()
    if font_config:
        config.update(font_config)
    
    font_name = config['font_name']
    font_size = config['body_size']
    
    # 后续行字段（2列）
    remaining_fields = ['试验机构', '试验环境', '试验标准', '试验条件', '规格要求', '试验数据', '试验结论']
    
    # 计算总行数
    total_rows = 2 + len(remaining_fields)
    
    # 创建表格（最多4列）
    table = doc.add_table(rows=total_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    
    # 设置列宽
    if col_widths and len(col_widths) == 4:
        for i, width in enumerate(col_widths):
            table.columns[i].width = Cm(width)
    
    # 第一行：开始日期 | 值 | 结束日期 | 值
    table.cell(0, 0).text = '开始日期'
    set_cell_font(table.cell(0, 0), font_name=font_name, font_size=font_size, bold=True, align_center=False)
    table.cell(0, 1).text = format_date_only(data_dict.get('开始日期', ''))
    set_cell_font(table.cell(0, 1), font_name=font_name, font_size=font_size, align_center=False)
    table.cell(0, 2).text = '结束日期'
    set_cell_font(table.cell(0, 2), font_name=font_name, font_size=font_size, bold=True, align_center=False)
    table.cell(0, 3).text = format_date_only(data_dict.get('结束日期', ''))
    set_cell_font(table.cell(0, 3), font_name=font_name, font_size=font_size, align_center=False)
    
    # 第二行：样机数量 | 值 | 样机编号 | 值
    table.cell(1, 0).text = '样机数量'
    set_cell_font(table.cell(1, 0), font_name=font_name, font_size=font_size, bold=True, align_center=False)
    table.cell(1, 1).text = format_quantity(data_dict.get('样机数量', ''))
    set_cell_font(table.cell(1, 1), font_name=font_name, font_size=font_size, align_center=False)
    table.cell(1, 2).text = '样机编号'
    set_cell_font(table.cell(1, 2), font_name=font_name, font_size=font_size, bold=True, align_center=False)
    table.cell(1, 3).text = format_sample_number(data_dict.get('样机编号', ''))
    set_cell_font(table.cell(1, 3), font_name=font_name, font_size=font_size, align_center=False)
    
    # 第三行起：字段名占1列，值合并3列
    for i, field in enumerate(remaining_fields):
        row_idx = i + 2
        
        # 合并第2-4列（值占3列）
        table.cell(row_idx, 1).merge(table.cell(row_idx, 2)).merge(table.cell(row_idx, 3))
        
        # 填充内容
        table.cell(row_idx, 0).text = field
        set_cell_font(table.cell(row_idx, 0), font_name=font_name, font_size=font_size, bold=True, align_center=False)
        value = data_dict.get(field, '')
        table.cell(row_idx, 1).text = str(value) if value else ''
        set_cell_font(table.cell(row_idx, 1), font_name=font_name, font_size=font_size, align_center=False)

    doc.add_paragraph()
    return table


class ExcelToWordReport:
    """Excel转Word报告主类"""

    # 默认字体配置
    DEFAULT_FONT_CONFIG = {
        'font_name': '微软雅黑',
        'title1_size': 16,      # 大标题：三号
        'title1_bold': True,
        'title2_size': 12,      # 次标题：小四
        'title2_bold': True,
        'title3_size': 12,      # 小标题：小四
        'title3_bold': False,
        'body_size': 10.5,      # 正文：五号
        'body_bold': False,
    }

    def __init__(self, excel_path, word_path=None, logo_path=None, report_number=None, 
                 company_name="公司", company_full_name=None, company_address=None, watermark_text=None, report_name=None, font_config=None,
                 testcase_config=None, table_widths=None):
        """
        初始化

        参数:
            excel_path: Excel文件路径
            word_path: Word输出路径，默认同名
            logo_path: Logo图片路径（页眉用）
            report_number: 报告编号（页眉用），默认自动生成
            company_name: 公司名称简称（页脚保密信息用）
            company_full_name: 公司名称全称（封面Logo下方显示）
            company_address: 公司地址（封面公司全称下方显示）
            watermark_text: 水印文字，如 "xxxx to xxxx"
            report_name: 报告名称（页眉、封面标题、文件名用），默认使用文件名
            font_config: 字体配置字典，可覆盖默认配置
            testcase_config: 小用例属性配置（优先级高于Excel），字典格式
            table_widths: 表格列宽配置字典，如 {'testcase': [2.5,4,2.5,4], 'summary': [2,5,6,2] }
        """
        self.excel_path = excel_path
        self.logo_path = logo_path
        self.report_number = report_number or self._generate_report_number()
        self.company_name = company_name
        self.company_full_name = company_full_name
        self.company_address = company_address
        self.watermark_text = watermark_text
        self.report_name = report_name
        
        # Word输出路径：优先使用report_name作为文件名
        if word_path:
            self.word_path = word_path
        elif report_name:
            # 使用report_name作为文件名
            output_dir = os.path.dirname(excel_path)
            self.word_path = os.path.join(output_dir, f'{report_name}.docx')
        else:
            self.word_path = os.path.splitext(excel_path)[0] + '_报告.docx'
        
        # 合并字体配置
        self.font_config = self.DEFAULT_FONT_CONFIG.copy()
        if font_config:
            self.font_config.update(font_config)
        
        # 小用例属性配置（优先级高于Excel）
        self.testcase_config = testcase_config or {}
        
        # 表格列宽配置
        self.table_widths = table_widths or {}

        # 读取Excel
        self.df = None
        self.excel_columns = []

        # 解析后的数据
        self.overview_data = {}
        self.big_cases = []
        self.summary_data = []
        self.col_name_to_idx = {}
    
    def _generate_report_number(self):
        """生成默认报告编号（日期+时间）"""
        return datetime.now().strftime("RPT%Y%m%d%H%M%S")

    def load_excel(self, sheet_name=0):
        """
        加载Excel文件（自动识别xlsx/xls格式）

        参数:
            sheet_name: sheet名称或索引
        """
        # 保存sheet名称，供后续使用
        self.sheet_name = sheet_name
        
        # 先检测实际格式
        actual_format = detect_excel_format(self.excel_path)
        file_ext = os.path.splitext(self.excel_path)[1].lower()
        
        # 根据实际格式选择引擎
        if actual_format == 'xls':
            engine = 'xlrd'
        else:
            engine = 'openpyxl'

        try:
            self.df = pd.read_excel(self.excel_path, sheet_name=sheet_name, header=None, engine=engine)
        except Exception as e:
            # 如果默认引擎失败，尝试另一个
            alt_engine = 'xlrd' if engine == 'openpyxl' else 'openpyxl'
            self.df = pd.read_excel(self.excel_path, sheet_name=sheet_name, header=None, engine=alt_engine)

        self.excel_columns = [str(col) for col in self.df.iloc[0].tolist() if pd.notna(col)]

    def find_test_project_column(self):
        """找到"试验项目"列（在所有行中搜索）"""
        for row_idx in range(len(self.df)):
            for col_idx, cell in enumerate(self.df.iloc[row_idx]):
                if pd.notna(cell) and '试验项目' in str(cell):
                    return col_idx, row_idx
        return None, None

    def find_merged_cells_info(self):
        """
        检测合并单元格（大用例）
        返回: [(起始行, 结束行, 大用例名), ...]
        """
        file_ext = os.path.splitext(self.excel_path)[1].lower()

        if file_ext == '.xls':
            return [], None, None

        # xlsx格式使用openpyxl读取合并单元格信息
        wb = load_workbook(self.excel_path)
        
        # 使用用户指定的sheet，而不是活动sheet
        if hasattr(self, 'sheet_name') and self.sheet_name is not None:
            if isinstance(self.sheet_name, int):
                ws = wb.worksheets[self.sheet_name]
            else:
                ws = wb[self.sheet_name]
        else:
            ws = wb.active

        # 找到试验项目列
        test_col = None
        header_row = None
        for row_idx in range(1, ws.max_row + 1):
            for col_idx in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=row_idx, column=col_idx).value
                if cell_value and '试验项目' in str(cell_value):
                    test_col = col_idx
                    header_row = row_idx
                    break
            if test_col:
                break

        if not test_col:
            return [], None, None

        merged_ranges = []
        header_merge_end = header_row
        
        # 找到标题行的合并范围
        for merged_range in ws.merged_cells.ranges:
            if merged_range.min_row == header_row and merged_range.min_col <= test_col <= merged_range.max_col:
                header_merge_end = merged_range.max_row
                break
        
        # 检测大用例：列合并（横向合并多列）且在标题行下方
        for merged_range in ws.merged_cells.ranges:
            if merged_range.min_col != merged_range.max_col and merged_range.min_row > header_merge_end:
                row_idx = merged_range.min_row
                cell_value = ws.cell(row=row_idx, column=merged_range.min_col).value
                if cell_value and str(cell_value).strip():
                    merged_ranges.append((row_idx, row_idx, str(cell_value).strip()))

        merged_ranges.sort(key=lambda x: x[0])
        
        # 验证大用例：必须下方有小用例
        valid_merged_ranges = []
        for i, (start_row, end_row, name) in enumerate(merged_ranges):
            next_start = merged_ranges[i + 1][0] if i + 1 < len(merged_ranges) else ws.max_row + 1
            
            has_small_case = False
            for row_idx in range(end_row + 1, next_start):
                cell_value = ws.cell(row=row_idx, column=test_col).value
                if cell_value and str(cell_value).strip():
                    has_small_case = True
                    break
            
            if has_small_case:
                valid_merged_ranges.append((start_row, end_row, name))
        
        return valid_merged_ranges, test_col, header_row

    def parse_overview_data(self, big_case_start_row):
        """
        解析概述数据（大用例之前的内容）

        参数:
            big_case_start_row: 第一个大用例的起始行
        """
        # 概述数据在大用例之前，通常是键值对形式
        for row_idx in range(1, big_case_start_row):
            row_data = self.df.iloc[row_idx]

            # 查找键值对（假设在D、E列或相邻列）
            for col_idx in range(len(row_data) - 1):
                key = row_data.iloc[col_idx]
                value = row_data.iloc[col_idx + 1]

                if pd.notna(key) and pd.notna(value):
                    key_str = str(key).strip()
                    value_str = str(value).strip()

                    # 智能匹配到概述字段
                    for field, keywords in OVERVIEW_FIELD_MAPPING.items():
                        if any(kw in key_str for kw in keywords):
                            self.overview_data[field] = value_str
                            break

    def parse_test_cases(self):
        """解析测试用例（大用例和小用例）"""
        merged_ranges, test_col, header_row = self.find_merged_cells_info()

        if not merged_ranges:
            self.parse_without_merge()
            return

        # 解析概述数据（第一个大用例之前）
        first_big_case_row = merged_ranges[0][0]
        self.parse_overview_data(first_big_case_row)

        # 建立列名到索引的映射
        self.build_column_mapping(header_row)

        # 解析大用例和小用例
        for i, (start_row, end_row, big_case_name) in enumerate(merged_ranges):
            big_case = {
                'name': big_case_name,
                'small_cases': []
            }

            next_start = merged_ranges[i + 1][0] if i + 1 < len(merged_ranges) else len(self.df) + 1

            for row_idx in range(end_row + 1, next_start):
                if row_idx <= len(self.df):
                    row_data = self.df.iloc[row_idx - 1]
                    small_case_name = row_data.iloc[test_col - 1] if test_col else None

                    if pd.notna(small_case_name) and str(small_case_name).strip():
                        small_case = {
                            'name': str(small_case_name),
                            'data': self.extract_testcase_data(row_data)
                        }
                        big_case['small_cases'].append(small_case)

                        self.summary_data.append({
                            '序号': len(self.summary_data) + 1,
                            '试验分类': clean_case_number(big_case_name),
                            '试验项目': clean_case_number(str(small_case_name)),
                            '测试结论': small_case['data'].get('试验结论', '')
                        })

            self.big_cases.append(big_case)

    def build_column_mapping(self, header_row):
        """
        建立列名到索引的映射（支持多行合并标题）

        参数:
            header_row: 标题行号（openpyxl格式，从1开始）
        """
        self.col_name_to_idx = {}
        
        # 根据文件格式选择方式
        file_ext = os.path.splitext(self.excel_path)[1].lower()
        
        if file_ext == '.xls':
            # xls格式直接从DataFrame读取
            if header_row and header_row < len(self.df):
                for col_idx, col_name in enumerate(self.df.iloc[header_row]):
                    if pd.notna(col_name):
                        self.col_name_to_idx[str(col_name).strip()] = col_idx
        else:
            # xlsx格式使用openpyxl读取标题行（支持多行合并标题）
            wb = load_workbook(self.excel_path)
            ws = wb.active
            
            # 先检测标题区域的合并单元格，找到标题行的结束行
            header_end_row = header_row
            for merged_range in ws.merged_cells.ranges:
                # 如果合并单元格跨越多行且包含标题行
                if (merged_range.min_row <= header_row <= merged_range.max_row and 
                    merged_range.min_row != merged_range.max_row):
                    header_end_row = max(header_end_row, merged_range.max_row)
            
            # 读取标题行（优先读取最后一行，因为子列名通常在下面）
            # 如果有多行标题，优先使用下面的行（子列名）
            for col_idx in range(1, ws.max_column + 1):
                # 从下往上找，优先取子列名
                cell_value = None
                for row in range(header_end_row, header_row - 1, -1):
                    val = ws.cell(row=row, column=col_idx).value
                    if val and str(val).strip():
                        cell_value = str(val).strip()
                        # 如果不是大类名称（如"样件"），就用这个
                        # 大类名称通常是合并单元格的父标题
                        break
                
                if cell_value:
                    self.col_name_to_idx[cell_value] = col_idx - 1

    def parse_without_merge(self):
        """无合并单元格时的备用解析"""
        test_col, header_row = self.find_test_project_column()
        if test_col is None:
            return

        self.build_column_mapping(header_row + 1 if header_row else 1)

        current_big_case = None

        for row_idx in range(1, len(self.df)):
            row_data = self.df.iloc[row_idx]
            test_project = row_data.iloc[test_col]

            if pd.notna(test_project):
                pass

    def extract_testcase_data(self, row_data):
        """
        从行数据中提取测试用例数据（配置优先，Excel补充）

        参数:
            row_data: DataFrame的一行

        返回:
            字典 {字段名: 值}
        """
        data = {}
        
        # 配置字段名 -> 内部字段名映射
        config_field_mapping = {
            'sample_quantity': '样机数量',
            'sample_number': '样机编号',
            'test_organization': '试验机构',
            'test_environment': '试验环境',
            'test_standard': '试验标准',
            'test_condition': '试验条件',
            'spec_requirement': '规格要求',
        }
        
        # 1. 先从配置读取（优先级最高）
        for config_key, field_name in config_field_mapping.items():
            config_value = self.testcase_config.get(config_key)
            if config_value:
                data[field_name] = config_value

        # 如果没有列名映射，尝试建立
        if not hasattr(self, 'col_name_to_idx') or not self.col_name_to_idx:
            self.col_name_to_idx = {}
            for idx, col_name in enumerate(self.df.iloc[0]):
                if pd.notna(col_name):
                    self.col_name_to_idx[str(col_name)] = idx

        # 2. 从Excel读取（补充配置没有的字段）
        for field, keywords in TESTCASE_FIELD_MAPPING.items():
            # 如果配置已经有值，跳过
            if field in data:
                continue
            
            for keyword in keywords:
                for col_name, idx in self.col_name_to_idx.items():
                    if keyword in col_name:
                        value = row_data.iloc[idx] if idx < len(row_data) else None
                        if pd.notna(value):
                            data[field] = str(value)
                            break
                if field in data:
                    break

        return data

    def generate_word_report(self):
        """生成Word报告"""
        doc = Document()

        # 禁用拼写和语法检查（避免红色下划线）
        doc.settings.spell_errors = False
        doc.settings.grammar_errors = False

        # 设置默认字体
        font_name = self.font_config.get('font_name', '微软雅黑')
        body_size = self.font_config.get('body_size', 10.5)
        
        doc.styles['Normal'].font.name = font_name
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        doc.styles['Normal'].font.size = Pt(body_size)
        
        # 设置标题样式（确保导航窗格能识别）
        setup_heading_styles(doc, self.font_config)

        # ===== 添加封面页 =====
        report_name = self.report_name or os.path.splitext(os.path.basename(self.word_path))[0]
        add_cover_page(doc, report_name, self.report_number, self.company_name, self.company_full_name, self.company_address, self.logo_path, self.font_config, self.overview_data)

        # ===== 添加声明页 =====
        add_declaration_page(doc, self.company_name, self.font_config)

        # ===== 添加页眉页脚 =====
        add_header_footer(doc, report_name, self.report_number, self.logo_path, self.company_name)

        # ===== 添加水印 =====
        if self.watermark_text:
            add_watermark_to_docx(doc, self.watermark_text)

        # ===== 目录 =====
        # 目录标题（居中，不作为标题样式）
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run('目录')
        toc_run.font.name = font_name
        toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        toc_run.font.size = Pt(self.font_config.get('title1_size', 16))
        toc_run.font.bold = self.font_config.get('title1_bold', True)
        toc_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 生成目录项（带超链接）
        toc_items = [
            ('1 概述', 'toc_1', 1),
            ('1.1 产品信息', 'toc_1_1', 2),
            ('1.2 试验信息', 'toc_1_2', 2),
            ('1.3 工作模式', 'toc_1_3', 2),
            ('1.4 测试仪器设备', 'toc_1_4', 2),
            ('2 试验结果汇总', 'toc_2', 1),
            ('3 测试数据', 'toc_3', 1),
        ]
        
        # 添加大用例和小用例到目录
        for big_idx, big_case in enumerate(self.big_cases, 1):
            clean_name = clean_case_number(big_case["name"])
            toc_items.append((f'3.{big_idx} {clean_name}', f'toc_3_{big_idx}', 2))
            for small_idx, small_case in enumerate(big_case['small_cases'], 1):
                clean_small_name = clean_case_number(small_case["name"])
                toc_items.append((f'3.{big_idx}.{small_idx} {clean_small_name}', f'toc_3_{big_idx}_{small_idx}', 3))
        
        # 输出目录（带超链接和页码格式）
        for item_text, bookmark_name, level in toc_items:
            toc_para = doc.add_paragraph()
            
            # 设置行间距为单倍行距（最紧凑）
            toc_para.paragraph_format.line_spacing = 1.0  # 单倍行距
            toc_para.paragraph_format.space_before = Pt(0)
            toc_para.paragraph_format.space_after = Pt(0)
            
            # 根据级别设置段落缩进和加粗
            if level == 1:
                toc_para.paragraph_format.left_indent = Cm(0)      # 大标题：无缩进
                is_bold = True
            elif level == 2:
                toc_para.paragraph_format.left_indent = Cm(0.3)    # 次标题/大用例：缩进0.3cm
                is_bold = True
            else:  # level == 3
                toc_para.paragraph_format.left_indent = Cm(0.6)    # 小用例：缩进0.6cm
                is_bold = False
            
            # 设置制表符：右对齐制表符在页面右侧，带点号前导符
            tab_stops = toc_para.paragraph_format.tab_stops
            # A4纸宽度21cm，左右边距各2.54cm，有效宽度约15.92cm
            tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            # 添加超链接标题（根据级别加粗）
            add_hyperlink(toc_para, item_text, bookmark_name, font_name, body_size, bold=is_bold)
            
            # 添加制表符（自动填充点号）
            tab_run = toc_para.add_run('\t')
            tab_run.font.name = font_name
            tab_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            tab_run.font.size = Pt(body_size)
            
            # 添加页码域（PAGEREF，动态引用书签页码）
            page_run = toc_para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = f' PAGEREF {bookmark_name} '
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            
            page_text = toc_para.add_run('1')
            page_text.font.name = font_name
            page_text._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            page_text.font.size = Pt(body_size)
            
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            page_run._r.append(fldChar1)
            page_run._r.append(instrText)
            page_run._r.append(fldChar2)
            page_run._r.append(page_text._r)
            page_run._r.append(fldChar3)
        
        doc.add_paragraph()  # 空行

        # ===== 1. 概述（大标题）=====
        h1 = add_heading_with_number(doc, '1 概述', level=1, font_config=self.font_config)
        add_bookmark(h1, 'toc_1')

        # 1.1 产品信息（次标题）
        h1_1 = add_heading_with_number(doc, '1.1 产品信息', level=2, font_config=self.font_config)
        add_bookmark(h1_1, 'toc_1_1')
        add_body_paragraph(doc, self.overview_data.get('产品信息', '（待填写）'), font_config=self.font_config)

        # 1.2 试验信息（次标题）
        h1_2 = add_heading_with_number(doc, '1.2 试验信息', level=2, font_config=self.font_config)
        add_bookmark(h1_2, 'toc_1_2')
        add_body_paragraph(doc, self.overview_data.get('试验信息', '（待填写）'), font_config=self.font_config)

        # 1.3 工作模式（次标题）
        h1_3 = add_heading_with_number(doc, '1.3 工作模式', level=2, font_config=self.font_config)
        add_bookmark(h1_3, 'toc_1_3')
        add_body_paragraph(doc, self.overview_data.get('工作模式', '（待填写）'), font_config=self.font_config)

        # 1.4 测试仪器设备（次标题）
        h1_4 = add_heading_with_number(doc, '1.4 测试仪器设备', level=2, font_config=self.font_config)
        add_bookmark(h1_4, 'toc_1_4')
        add_body_paragraph(doc, self.overview_data.get('测试仪器设备', '（待填写）'), font_config=self.font_config)

        # ===== 2. 试验结果汇总（大标题）=====
        h2 = add_heading_with_number(doc, '2 试验结果汇总', level=1, font_config=self.font_config)
        add_bookmark(h2, 'toc_2')

        # 创建汇总表格
        summary_table = doc.add_table(rows=len(self.summary_data) + 1, cols=4)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_border(summary_table)
        
        # 设置列宽
        summary_widths = self.table_widths.get('summary', [2.67, 6.4, 5.75, 2.67])
        for i, width in enumerate(summary_widths):
            summary_table.columns[i].width = Cm(width)

        # 表头
        headers = ['序号', '试验分类', '试验项目', '测试结论']
        for i, header in enumerate(headers):
            cell = summary_table.cell(0, i)
            cell.text = header
            set_cell_font(cell, font_name=font_name, font_size=body_size, bold=True)

        # 数据行
        for row_idx, item in enumerate(self.summary_data):
            for col_idx, header in enumerate(headers):
                cell = summary_table.cell(row_idx + 1, col_idx)
                cell.text = str(item.get(header, ''))
                set_cell_font(cell, font_name=font_name, font_size=body_size)
        
        # 调试输出
        print("\n试验结果汇总数据:")
        for idx, item in enumerate(self.summary_data):
            print(f"  行{idx+1}: 试验分类={item.get('试验分类', '')}, 试验项目={item.get('试验项目', '')}")
        
        # 合并相同试验分类的单元格并设置居中
        if len(self.summary_data) > 0:
            # 找出所有需要合并的范围
            merge_ranges = []  # [(start_row, end_row), ...]
            current_category = None
            merge_start = 1
            
            for row_idx, item in enumerate(self.summary_data, 1):
                category = item.get('试验分类', '')
                
                if category != current_category:
                    if current_category is not None:
                        merge_ranges.append((merge_start, row_idx - 1))
                    current_category = category
                    merge_start = row_idx
            
            # 添加最后一个分类
            if current_category is not None:
                merge_ranges.append((merge_start, len(self.summary_data)))
            
            # 执行合并
            print(f"\n合并范围: {merge_ranges}")
            for start, end in merge_ranges:
                if end > start:  # 至少2行才合并
                    print(f"  合并第{start}行到第{end}行")
                    
                    # 获取第一个单元格的值
                    first_cell = summary_table.cell(start, 1)
                    value = first_cell.text
                    
                    # 清空其他单元格的内容
                    for row_idx in range(start + 1, end + 1):
                        cell = summary_table.cell(row_idx, 1)
                        for para in cell.paragraphs:
                            para.clear()
                    
                    # 执行合并
                    merged_cell = first_cell.merge(summary_table.cell(end, 1))
                    
                    # 重新设置值（合并后可能丢失）
                    # 清空合并后的段落，重新添加一个居中的段落
                    for para in merged_cell.paragraphs:
                        para.clear()
                    
                    # 添加新的段落并设置居中
                    new_para = merged_cell.paragraphs[0]
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = new_para.add_run(value)
                    run.font.name = font_name
                    run.font.size = Pt(body_size)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    
                    # 设置垂直居中
                    merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        doc.add_paragraph()

        # ===== 3. 测试数据 =====
        h3 = add_heading_with_number(doc, '3 测试数据', level=1, font_config=self.font_config)
        add_bookmark(h3, 'toc_3')
        
        # 验证大用例章节位置
        print("\n验证大用例章节位置...")
        validation_results = validate_chapter_position(self.big_cases)
        invalid_count = sum(1 for r in validation_results if not r['valid'])
        if invalid_count > 0:
            print(f"⚠️ 发现 {invalid_count} 个大用例位置不匹配")
        else:
            print("✓ 所有大用例位置正确")

        for big_idx, big_case in enumerate(self.big_cases, 1):
            clean_name = clean_case_number(big_case["name"])
            h3_big = add_heading_with_number(doc, f'3.{big_idx} {clean_name}', level=2, font_config=self.font_config)
            add_bookmark(h3_big, f'toc_3_{big_idx}')

            for small_idx, small_case in enumerate(big_case['small_cases'], 1):
                clean_small_name = clean_case_number(small_case["name"])
                h3_small = add_heading_with_number(doc, f'3.{big_idx}.{small_idx} {clean_small_name}', level=3, font_config=self.font_config)
                add_bookmark(h3_small, f'toc_3_{big_idx}_{small_idx}')
                create_testcase_table(doc, small_case['data'], font_config=self.font_config, 
                                      col_widths=self.table_widths.get('testcase'))

        # 保存文档
        # 先检查并关闭占用该文件的Word文档
        close_word_document(self.word_path)
        
        # 如果文件已存在，删除旧文件
        if os.path.exists(self.word_path):
            try:
                os.remove(self.word_path)
                print(f"  已删除旧文件: {os.path.basename(self.word_path)}")
            except Exception as e:
                print(f"  警告: 无法删除旧文件 - {e}")
        
        try:
            doc.save(self.word_path)
            print(f"Word报告已生成: {self.word_path}")
        except PermissionError:
            print(f"\n❌ 错误: 文件仍被占用，无法保存!")
            print(f"   请手动关闭 Word 中打开的 '{os.path.basename(self.word_path)}' 后重试")
            return None
        except Exception as e:
            print(f"\n❌ 保存失败: {e}")
            return None
        
        # 自动更新目录页码
        update_toc_in_word(self.word_path)
        
        return self.word_path


def detect_excel_format(file_path):
    """
    检测Excel文件的实际格式（通过文件头）
    
    返回:
        'xlsx' - Excel 2007+ (Office Open XML)
        'xls' - Excel 97-2003 (BIFF)
    """
    with open(file_path, 'rb') as f:
        header = f.read(8)
    
    # xlsx 实际是 ZIP 压缩包，以 PK 开头
    if header[:2] == b'PK':
        return 'xlsx'
    
    # xls 是 OLE 复合文档，以 D0 CF 11 E0 开头
    if header[:4] == b'\xD0\xCF\x11\xE0':
        return 'xls'
    
    # 默认返回xlsx
    return 'xlsx'


def list_sheets(excel_path):
    """
    列出Excel中所有sheet名称

    参数:
        excel_path: Excel文件路径

    返回:
        sheet名称列表
    """
    # 先检测实际格式
    actual_format = detect_excel_format(excel_path)
    file_ext = os.path.splitext(excel_path)[1].lower()
    
    print(f"文件扩展名: {file_ext}, 实际格式: {actual_format}")
    
    if actual_format == 'xls':
        # xls格式用xlrd
        try:
            wb = xlrd.open_workbook(excel_path)
            return wb.sheet_names()
        except Exception as e:
            print(f"xlrd读取失败: {e}")
            # 尝试用pandas
            xl = pd.ExcelFile(excel_path, engine='xlrd')
            return xl.sheet_names
    else:
        # xlsx格式用openpyxl
        try:
            wb = load_workbook(excel_path, read_only=True)
            sheets = wb.sheetnames
            wb.close()
            return sheets
        except Exception as e:
            print(f"openpyxl读取失败: {e}")
            # 可能实际是xls格式，尝试xlrd
            try:
                wb = xlrd.open_workbook(excel_path)
                return wb.sheet_names()
            except:
                pass
            # 最后尝试pandas
            xl = pd.ExcelFile(excel_path)
            return xl.sheet_names


def process_sheets(excel_path, sheets=None, output_dir=None, merge=False, 
                   logo_path=None, report_number=None, company_name="公司", company_full_name=None, company_address=None,
                   watermark_text=None, report_name=None, font_config=None,
                   testcase_config=None, table_widths=None):
    """
    处理指定的sheet，生成Word报告

    参数:
        excel_path: Excel文件路径
        sheets: 要处理的sheet列表，可以是：
                - None: 处理所有sheet
                - int: 单个sheet索引
                - str: 单个sheet名称
                - list: [0, 1, 2] 或 ["Sheet1", "Sheet2"]
        output_dir: 输出目录，None表示与Excel同目录
        merge: 是否合并多个sheet到一个Word文件
        logo_path: Logo图片路径
        report_number: 报告编号
        company_name: 公司名称简称
        company_full_name: 公司名称全称
        company_address: 公司地址
        watermark_text: 水印文字
        report_name: 报告名称（页眉用）
        font_config: 字体配置字典
        testcase_config: 小用例属性配置（优先级高于Excel）
        table_widths: 表格列宽配置字典

    返回:
        生成的Word文件路径列表
    """
    # 获取所有sheet名称
    all_sheets = list_sheets(excel_path)
    print(f"Excel包含 {len(all_sheets)} 个sheet: {all_sheets}")

    # 确定要处理的sheet
    sheets_to_process = _resolve_sheets(all_sheets, sheets)
    
    if not sheets_to_process:
        return []

    print(f"将处理 {len(sheets_to_process)} 个sheet: {sheets_to_process}")

    # 确定输出目录
    if output_dir is None:
        output_dir = os.path.dirname(excel_path)

    # 合并模式
    if merge and len(sheets_to_process) > 1:
        return _merge_sheets_to_word(excel_path, sheets_to_process, output_dir, 
                                     logo_path, report_number, company_name, company_full_name, company_address,
                                     watermark_text, report_name, font_config,
                                     testcase_config, table_widths)
    
    # 单独生成模式
    return _generate_separate_reports(excel_path, sheets_to_process, output_dir,
                                      logo_path, report_number, company_name, company_full_name, company_address,
                                      watermark_text, report_name, font_config,
                                      testcase_config, table_widths)


def _resolve_sheets(all_sheets, sheets):
    """解析要处理的sheet列表"""
    sheets_to_process = []

    if sheets is None:
        sheets_to_process = all_sheets
    elif isinstance(sheets, int):
        if 0 <= sheets < len(all_sheets):
            sheets_to_process = [all_sheets[sheets]]
        else:
            print(f"错误: sheet索引 {sheets} 超出范围 (0-{len(all_sheets)-1})")
    elif isinstance(sheets, str):
        if sheets in all_sheets:
            sheets_to_process = [sheets]
        else:
            print(f"错误: 未找到名为 '{sheets}' 的sheet")
    elif isinstance(sheets, list):
        for s in sheets:
            if isinstance(s, int):
                if 0 <= s < len(all_sheets):
                    sheets_to_process.append(all_sheets[s])
                else:
                    print(f"警告: sheet索引 {s} 超出范围，跳过")
            elif isinstance(s, str):
                if s in all_sheets:
                    sheets_to_process.append(s)
                else:
                    print(f"警告: 未找到名为 '{s}' 的sheet，跳过")
    else:
        print(f"错误: 不支持的sheets参数类型: {type(sheets)}")

    return list(dict.fromkeys(sheets_to_process))


def _generate_separate_reports(excel_path, sheets_to_process, output_dir,
                                logo_path=None, report_number=None, company_name="公司", company_full_name=None, company_address=None,
                                watermark_text=None, report_name=None, font_config=None,
                                testcase_config=None, table_widths=None):
    """为每个sheet生成单独的Word报告"""
    output_files = []
    base_name = os.path.splitext(os.path.basename(excel_path))[0]

    for sheet_name in sheets_to_process:
        print(f"\n{'='*50}")
        print(f"正在处理sheet: {sheet_name}")
        print('='*50)

        # 生成输出文件名：优先使用report_name
        if report_name:
            word_path = os.path.join(output_dir, f"{report_name}.docx")
        elif len(sheets_to_process) == 1:
            word_path = os.path.join(output_dir, f"{base_name}_报告.docx")
        else:
            word_path = os.path.join(output_dir, f"{base_name}_{sheet_name}_报告.docx")

        try:
            converter = ExcelToWordReport(excel_path, word_path, logo_path, report_number, 
                                          company_name, company_full_name, company_address, watermark_text, report_name, font_config,
                                          testcase_config, table_widths)
            converter.load_excel(sheet_name)
            converter.parse_test_cases()
            output_path = converter.generate_word_report()
            output_files.append(output_path)
        except Exception as e:
            print(f"处理sheet '{sheet_name}' 时出错: {e}")
            traceback.print_exc()

    return output_files


def _merge_sheets_to_word(excel_path, sheets_to_process, output_dir,
                          logo_path=None, report_number=None, company_name="公司", company_full_name=None, company_address=None,
                          watermark_text=None, report_name=None, font_config=None,
                          testcase_config=None, table_widths=None):
    """将多个sheet合并到一个Word文件"""
    base_name = os.path.splitext(os.path.basename(excel_path))[0]
    word_path = os.path.join(output_dir, f"{base_name}_合并报告.docx")
    
    # 获取字体配置
    fc = font_config or {}
    font_name = fc.get('font_name', '微软雅黑')
    body_size = fc.get('body_size', 10.5)
    
    # 创建合并文档
    doc = Document()
    doc.styles['Normal'].font.name = font_name
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    doc.styles['Normal'].font.size = Pt(body_size)
    
    # 添加页眉页脚
    actual_report_name = report_name or f"{base_name}_合并报告"
    actual_report_number = report_number or f"RPT{datetime.now().strftime('%Y%m%d%H%M%S')}"
    add_header_footer(doc, actual_report_name, actual_report_number, logo_path, company_name)
    
    # 添加水印
    if watermark_text:
        add_watermark_to_docx(doc, watermark_text)
    
    # 添加总标题
    title = doc.add_heading(f'{base_name} 测试报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    print(f"\n{'='*50}")
    print(f"合并模式: 将 {len(sheets_to_process)} 个sheet合并到一个Word")
    print('='*50)
    
    for idx, sheet_name in enumerate(sheets_to_process):
        print(f"\n正在处理sheet [{idx+1}/{len(sheets_to_process)}]: {sheet_name}")
        
        try:
            # 创建临时转换器解析数据
            converter = ExcelToWordReport(excel_path)
            converter.load_excel(sheet_name)
            converter.parse_test_cases()
            
            # 添加分页符（第一个sheet不分页）
            if idx > 0:
                doc.add_page_break()
            
            # 添加sheet标题
            doc.add_heading(f'{idx+1} {sheet_name}', level=1)
            
            # 添加概述
            doc.add_heading(f'{idx+1}.1 概述', level=2)
            for field in ['产品信息', '试验信息', '工作模式', '测试仪器设备']:
                doc.add_paragraph(f"{field}: {converter.overview_data.get(field, '（待填写）')}")
            
            # 添加试验结果汇总
            if converter.summary_data:
                doc.add_heading(f'{idx+1}.2 试验结果汇总', level=2)
                summary_table = doc.add_table(rows=len(converter.summary_data) + 1, cols=4)
                summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_border(summary_table)
                
                # 设置列宽
                summary_widths = table_widths.get('summary', [2.67, 6.4, 5.75, 2.67]) if table_widths else [2.67, 6.4, 5.75, 2.67]
                for i, width in enumerate(summary_widths):
                    summary_table.columns[i].width = Cm(width)
                
                headers = ['序号', '试验分类', '试验项目', '测试结论']
                for i, header in enumerate(headers):
                    cell = summary_table.cell(0, i)
                    cell.text = header
                    set_cell_font(cell, bold=True)
                
                for row_idx, item in enumerate(converter.summary_data):
                    for col_idx, header in enumerate(headers):
                        cell = summary_table.cell(row_idx + 1, col_idx)
                        cell.text = str(item.get(header, ''))
                        set_cell_font(cell)
                
                # 合并相同试验分类的单元格并居中
                if len(converter.summary_data) > 1:
                    merge_ranges = []
                    current_category = None
                    merge_start = 1
                    
                    for row_idx, item in enumerate(converter.summary_data, 1):
                        category = item.get('试验分类', '')
                        
                        if category != current_category:
                            if current_category is not None:
                                merge_ranges.append((merge_start, row_idx - 1))
                            current_category = category
                            merge_start = row_idx
                    
                    if current_category is not None:
                        merge_ranges.append((merge_start, len(converter.summary_data)))
                    
                    for start, end in merge_ranges:
                        if end > start:
                            first_cell = summary_table.cell(start, 1)
                            value = first_cell.text
                            
                            for row_idx in range(start + 1, end + 1):
                                cell = summary_table.cell(row_idx, 1)
                                for para in cell.paragraphs:
                                    para.clear()
                            
                            merged_cell = first_cell.merge(summary_table.cell(end, 1))
                            
                            for para in merged_cell.paragraphs:
                                para.clear()
                            
                            new_para = merged_cell.paragraphs[0]
                            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = new_para.add_run(value)
                            run.font.name = font_name
                            run.font.size = Pt(body_size)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                            
                            merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # 添加测试数据
            if converter.big_cases:
                doc.add_heading(f'{idx+1}.3 测试数据', level=2)
                for big_idx, big_case in enumerate(converter.big_cases):
                    clean_big_name = clean_case_number(big_case["name"])
                    doc.add_heading(f'{idx+1}.3.{big_idx+1} {clean_big_name}', level=3)
                    for small_idx, small_case in enumerate(big_case['small_cases']):
                        clean_small_name = clean_case_number(small_case["name"])
                        doc.add_heading(f'{idx+1}.3.{big_idx+1}.{small_idx+1} {clean_small_name}', level=4)
                        create_testcase_table(doc, small_case['data'], col_widths=table_widths.get('testcase') if table_widths else None)
            
            print(f"  ✓ {sheet_name} 处理完成")
            
        except Exception as e:
            print(f"  ✗ 处理sheet '{sheet_name}' 时出错: {e}")
            traceback.print_exc()
    
    # 保存文档
    # 先检查并关闭占用该文件的Word文档
    close_word_document(word_path)
    
    # 如果文件已存在，删除旧文件
    if os.path.exists(word_path):
        try:
            os.remove(word_path)
            print(f"  已删除旧文件: {os.path.basename(word_path)}")
        except Exception as e:
            print(f"  警告: 无法删除旧文件 - {e}")
    
    try:
        doc.save(word_path)
        print(f"\n{'='*50}")
        print(f"合并报告已生成: {word_path}")
    except PermissionError:
        print(f"\n❌ 错误: 文件仍被占用，无法保存!")
        print(f"   请手动关闭 Word 中打开的 '{os.path.basename(word_path)}' 后重试")
        return []
    except Exception as e:
        print(f"\n❌ 保存失败: {e}")
        return []
    
    return [word_path]


def load_config(config_path):
    """
    从配置文件读取参数
    
    配置文件格式（每行一个配置，#开头为注释）：
    excel_file=D:\AI\test_data.xlsx
    logo_path=D:\AI\logo.png
    font_name=微软雅黑
    title1_size=16
    
    参数:
        config_path: 配置文件路径
        
    返回:
        配置字典
    """
    config = {
        'excel_file': r"D:\AI\test_data.xlsx",
        'output_dir': None,
        'logo_path': None,
        'report_number': None,
        'company_name': "公司",
        'report_name': None,
        'watermark_text': None,
        'sheets': None,
        'merge': False,
        # 字体配置
        'font_name': '微软雅黑',
        'title1_size': 16,
        'title1_bold': True,
        'title2_size': 12,
        'title2_bold': True,
        'title3_size': 12,
        'title3_bold': False,
        'body_size': 10.5,
        'body_bold': False,
        # 小用例属性配置（优先级高于Excel）
        'sample_quantity': None,      # 样机数量
        'sample_number': None,        # 样机编号
        'test_organization': None,    # 试验机构
        'test_environment': None,     # 试验环境
        'test_standard': None,        # 试验标准
        'test_condition': None,       # 试验条件/试验方法
        'spec_requirement': None,     # 规格要求
        # 中文配置项（兼容）
        '样机数量': None,
        '样机编号': None,
        '试验机构': None,
        '试验环境': None,
        '试验标准': None,
        '试验条件': None,
        '规格要求': None,
        # 表格列宽配置（单位：厘米）
        'testcase_table_widths': None,    # 小用例属性表格列宽
        'summary_table_widths': None,     # 试验结果汇总表格列宽
    }
    
    if not os.path.exists(config_path):
        print(f"配置文件不存在: {config_path}")
        print("将使用默认配置")
        return config
    
    with open(config_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            
            # 跳过空行和注释
            if not line or line.startswith('#'):
                continue
            
            # 解析 key=value
            if '=' in line:
                key, value = line.split('=', 1)
                key = key.strip()
                value = value.strip()
                
                # 处理特殊值
                if value.lower() == 'none' or value == '':
                    config[key] = None
                elif value.lower() == 'true':
                    config[key] = True
                elif value.lower() == 'false':
                    config[key] = False
                elif key in ['title1_size', 'title2_size', 'title3_size', 'body_size']:
                    # 字体大小配置（支持小数）
                    try:
                        config[key] = float(value)
                    except:
                        config[key] = float(config.get(key, 10.5))
                elif key == 'sheets':
                    if ',' in value:
                        config[key] = [int(s.strip()) if s.strip().isdigit() else s.strip() for s in value.split(',')]
                    elif value.isdigit():
                        config[key] = int(value)
                    else:
                        config[key] = value
                else:
                    # 支持 \n 换行符
                    config[key] = value.replace('\\n', '\n')
    
    print(f"已从配置文件加载: {config_path}")
    return config


def create_default_config(config_path):
    """
    创建默认配置文件
    
    参数:
        config_path: 配置文件路径
    """
    default_config = '''# Excel转Word报告配置文件
# 每行一个配置，格式：key=value
# #开头的行为注释，会被忽略

# ===== 必填项 =====
# Excel输入文件路径
excel_file=D:\\AI\\test_data.xlsx

# ===== 页眉页脚配置 =====
# Logo图片路径（可选，None则不显示）
logo_path=None

# 报告编号（可选，None则自动生成）
report_number=None

# 公司名称简称（用于页脚保密信息）
company_name=公司

# 公司名称全称（用于封面Logo下方显示）
company_full_name=None

# 公司地址（用于封面公司名称下方显示）
company_address=None

# 报告名称（页眉显示，None则使用文件名）
report_name=None

# ===== 水印配置 =====
# 水印文字（设置后会在页面中央显示灰色斜向水印）
# 示例：watermark_text=张三 to 李四
# 不需要水印则设为None
watermark_text=None

# ===== Sheet配置 =====
# 要处理的sheet（None=全部，数字=索引，字符串=名称，逗号分隔=多个）
sheets=None

# 合并模式（True=合并到一个Word，False=每个sheet单独生成）
merge=False

# ===== 字体配置 =====
# 字体名称（默认微软雅黑）
font_name=微软雅黑

# 大标题字号（默认16=三号）
title1_size=16
# 大标题是否加粗（true/false）
title1_bold=true

# 次标题字号（默认12=小四）
title2_size=12
# 次标题是否加粗（true/false）
title2_bold=true

# 小标题字号（默认12=小四）
title3_size=12
# 小标题是否加粗（true/false）
title3_bold=false

# 正文字号（默认10.5=五号）
body_size=10.5
# 正文是否加粗（true/false）
body_bold=false

# ===== 输出配置 =====
# 输出目录（None表示与Excel同目录）
output_dir=None
'''
    
    with open(config_path, 'w', encoding='utf-8') as f:
        f.write(default_config)
    
    print(f"已创建默认配置文件: {config_path}")


def main():
    """主函数"""
    # ===== 配置文件路径 =====
    config_path = r"D:\AI\report_config.txt"
    
    # 如果配置文件不存在，创建默认配置
    if not os.path.exists(config_path):
        create_default_config(config_path)
        print("\n请修改配置文件后重新运行：")
        print(f"  {config_path}")
        return
    
    # 从配置文件加载参数
    config = load_config(config_path)
    
    # 提取字体配置
    font_config = {
        'font_name': config.get('font_name', '微软雅黑'),
        'title1_size': config.get('title1_size', 16),
        'title1_bold': config.get('title1_bold', True),
        'title2_size': config.get('title2_size', 12),
        'title2_bold': config.get('title2_bold', True),
        'title3_size': config.get('title3_size', 12),
        'title3_bold': config.get('title3_bold', False),
        'body_size': config.get('body_size', 10.5),
        'body_bold': config.get('body_bold', False),
    }
    
    # 提取小用例属性配置（支持中文和英文配置项，中文优先）
    testcase_config = {
        'sample_quantity': config.get('样机数量') or config.get('sample_quantity'),
        'sample_number': config.get('样机编号') or config.get('sample_number'),
        'test_organization': config.get('试验机构') or config.get('test_organization'),
        'test_environment': config.get('试验环境') or config.get('test_environment'),
        'test_standard': config.get('试验标准') or config.get('test_standard'),
        'test_condition': config.get('试验条件') or config.get('test_condition'),
        'spec_requirement': config.get('规格要求') or config.get('spec_requirement'),
    }
    
    # 打印配置信息
    print(f"\n小用例属性配置:")
    for key, value in testcase_config.items():
        if value:
            print(f"  {key}: {value}")
    
    # 解析表格列宽配置
    table_widths = {}
    
    # 小用例属性表格列宽
    testcase_widths_str = config.get('testcase_table_widths') or config.get('小用例表格列宽')
    if testcase_widths_str:
        try:
            table_widths['testcase'] = [float(w.strip()) for w in testcase_widths_str.split(',')]
            print(f"  小用例表格列宽: {table_widths['testcase']} cm")
        except:
            print(f"  警告: 小用例表格列宽格式错误，使用默认值")
    
    # 试验结果汇总表格列宽
    summary_widths_str = config.get('summary_table_widths') or config.get('汇总表格列宽')
    if summary_widths_str:
        try:
            table_widths['summary'] = [float(w.strip()) for w in summary_widths_str.split(',')]
            print(f"  汇总表格列宽: {table_widths['summary']} cm")
        except:
            print(f"  警告: 汇总表格列宽格式错误，使用默认值")
    
    # 检查文件是否存在
    if not os.path.exists(config['excel_file']):
        print(f"错误: Excel文件不存在 - {config['excel_file']}")
        print("请修改配置文件中的 excel_file")
        return

    # 先列出所有sheet
    all_sheets = list_sheets(config['excel_file'])
    print(f"\n可用sheet列表:")
    for i, name in enumerate(all_sheets):
        print(f"  [{i}] {name}")

    # 处理sheet
    output_files = process_sheets(
        config['excel_file'], 
        config['sheets'], 
        config['output_dir'], 
        config['merge'], 
        logo_path=config['logo_path'], 
        report_number=config['report_number'], 
        company_name=config['company_name'],
        company_full_name=config.get('company_full_name'),
        company_address=config.get('company_address'),
        watermark_text=config['watermark_text'],
        report_name=config['report_name'],
        font_config=font_config,
        testcase_config=testcase_config,
        table_widths=table_widths
    )

    # 输出结果
    print(f"\n{'='*50}")
    print(f"转换完成！共生成 {len(output_files)} 个报告:")
    for f in output_files:
        print(f"  - {f}")


if __name__ == '__main__':
    main()
