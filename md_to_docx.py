# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_report():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('GB 44495-2024 电驱MCU信息安全测试报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 文档信息
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('文档编号：', '[项目编号]-SEC-TEST-[年份]-[序号]'),
        ('版本：', 'V1.0'),
        ('密级：', '内部/机密'),
        ('编制日期：', '2026-XX-XX')
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # 文档修订历史
    doc.add_heading('文档修订历史', level=1)
    rev_table = doc.add_table(rows=2, cols=4)
    rev_table.style = 'Table Grid'
    headers = ['版本', '日期', '修订人', '修订内容']
    for i, h in enumerate(headers):
        cell = rev_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    rev_table.rows[1].cells[0].text = 'V1.0'
    rev_table.rows[1].cells[1].text = '2026-XX-XX'
    rev_table.rows[1].cells[2].text = 'XXX'
    rev_table.rows[1].cells[3].text = '初始版本'
    
    doc.add_page_break()
    
    # 目录
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 概述',
        '2. 测试样品信息',
        '3. 测试人员清单',
        '4. 测试设备清单',
        '5. 测试结论',
        '6. 问题清单',
        '7. 测试详情'
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # 1. 概述
    doc.add_heading('1. 概述', level=1)
    
    doc.add_heading('1.1 测试目的', level=2)
    doc.add_paragraph('本次测试依据 GB 44495-2024《汽车整车信息安全技术要求》国家标准，对新能源汽车电驱系统电机控制器（MCU）进行信息安全测试，验证其是否满足国家标准要求，识别潜在安全风险，为产品信息安全合规性提供依据。')
    
    doc.add_heading('1.2 测试依据', level=2)
    basis_table = doc.add_table(rows=6, cols=3)
    basis_table.style = 'Table Grid'
    basis_headers = ['序号', '依据文件', '版本/编号']
    for i, h in enumerate(basis_headers):
        cell = basis_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    basis_data = [
        ('1', 'GB 44495-2024《汽车整车信息安全技术要求》', '2024版'),
        ('2', 'GB/T 40857-2021《汽车信息安全通用技术要求》', '-'),
        ('3', 'ISO/SAE 21434《道路车辆-网络安全工程》', '2021版'),
        ('4', 'UN R155《网络安全和网络安全管理系统》', '-'),
        ('5', '企业内部信息安全测试规范', '[企业规范编号]')
    ]
    for i, row_data in enumerate(basis_data, 1):
        for j, val in enumerate(row_data):
            basis_table.rows[i].cells[j].text = val
    
    doc.add_heading('1.3 测试范围', level=2)
    doc.add_paragraph('本次测试覆盖电驱MCU以下方面：')
    scope_items = [
        '车载软件升级安全',
        '车载通信安全',
        '外部接口安全',
        '数据安全',
        '硬件安全',
        '软件安全'
    ]
    for item in scope_items:
        doc.add_paragraph(f'• {item}')
    
    doc.add_heading('1.4 测试环境', level=2)
    env_table = doc.add_table(rows=5, cols=2)
    env_table.style = 'Table Grid'
    env_data = [
        ('项目', '描述'),
        ('测试地点', '[实验室名称/地点]'),
        ('测试周期', '2026年XX月XX日 - 2026年XX月XX日'),
        ('测试环境温度', '23±5℃'),
        ('测试环境湿度', '45%-75% RH')
    ]
    for i, (label, value) in enumerate(env_data):
        env_table.rows[i].cells[0].text = label
        env_table.rows[i].cells[1].text = value
        if i == 0:
            set_cell_shading(env_table.rows[i].cells[0], 'D9E2F3')
            set_cell_shading(env_table.rows[i].cells[1], 'D9E2F3')
    
    doc.add_heading('1.5 缩略语说明', level=2)
    abbr_table = doc.add_table(rows=11, cols=3)
    abbr_table.style = 'Table Grid'
    abbr_headers = ['缩略语', '全称', '中文含义']
    for i, h in enumerate(abbr_headers):
        cell = abbr_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    abbr_data = [
        ('MCU', 'Motor Control Unit', '电机控制器'),
        ('ECU', 'Electronic Control Unit', '电子控制单元'),
        ('CAN', 'Controller Area Network', '控制器局域网'),
        ('CAN-FD', 'CAN with Flexible Data-rate', '灵活数据速率CAN'),
        ('UDS', 'Unified Diagnostic Services', '统一诊断服务'),
        ('OTA', 'Over-The-Air', '空中下载技术'),
        ('HSM', 'Hardware Security Module', '硬件安全模块'),
        ('SE', 'Secure Element', '安全元件'),
        ('TARA', 'Threat Analysis and Risk Assessment', '威胁分析与风险评估'),
        ('SecOC', 'Secure Onboard Communication', '安全车载通信')
    ]
    for i, row_data in enumerate(abbr_data, 1):
        for j, val in enumerate(row_data):
            abbr_table.rows[i].cells[j].text = val
    
    doc.add_page_break()
    
    # 2. 测试样品信息
    doc.add_heading('2. 测试样品信息', level=1)
    
    doc.add_heading('2.1 样品基本信息', level=2)
    sample_table = doc.add_table(rows=12, cols=2)
    sample_table.style = 'Table Grid'
    sample_data = [
        ('项目', '内容'),
        ('样品名称', '电机控制器（MCU）'),
        ('样品型号', '[型号]'),
        ('样品编号', '[序列号]'),
        ('硬件版本', '[版本号]'),
        ('软件版本', '[版本号]'),
        ('固件版本', '[版本号]'),
        ('bootloader版本', '[版本号]'),
        ('生产厂商', '[厂商名称]'),
        ('生产日期', '[日期]'),
        ('样品数量', '[数量]台'),
        ('样品状态', '全新/正常使用/其他')
    ]
    for i, (label, value) in enumerate(sample_data):
        sample_table.rows[i].cells[0].text = label
        sample_table.rows[i].cells[1].text = value
        if i == 0:
            set_cell_shading(sample_table.rows[i].cells[0], 'D9E2F3')
            set_cell_shading(sample_table.rows[i].cells[1], 'D9E2F3')
    
    doc.add_heading('2.2 样品技术参数', level=2)
    param_table = doc.add_table(rows=12, cols=2)
    param_table.style = 'Table Grid'
    param_data = [
        ('项目', '参数'),
        ('额定电压', '[XX]V'),
        ('额定功率', '[XX]kW'),
        ('峰值功率', '[XX]kW'),
        ('控制芯片型号', '[型号]'),
        ('主频', '[XX]MHz'),
        ('Flash容量', '[XX]KB'),
        ('RAM容量', '[XX]KB'),
        ('通信接口', 'CAN/CAN-FD/LIN/Ethernet等'),
        ('是否集成HSM', '是/否'),
        ('HSM型号（如适用）', '[型号]')
    ]
    for i, (label, value) in enumerate(param_data):
        param_table.rows[i].cells[0].text = label
        param_table.rows[i].cells[1].text = value
        if i == 0:
            set_cell_shading(param_table.rows[i].cells[0], 'D9E2F3')
            set_cell_shading(param_table.rows[i].cells[1], 'D9E2F3')
    
    doc.add_heading('2.3 样品接口信息', level=2)
    interface_table = doc.add_table(rows=8, cols=3)
    interface_table.style = 'Table Grid'
    interface_headers = ['接口类型', '数量', '用途']
    for i, h in enumerate(interface_headers):
        cell = interface_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    interface_data = [
        ('CAN总线接口', '[X]路', '整车CAN通信'),
        ('CAN-FD接口', '[X]路', '高速数据通信'),
        ('Ethernet接口', '[X]路', '诊断/刷写'),
        ('LIN接口', '[X]路', '从设备通信'),
        ('调试接口(JTAG/SWD)', '[X]路', '开发调试'),
        ('串口(UART)', '[X]路', '调试/日志'),
        ('OBD接口', '[X]路', '诊断接口')
    ]
    for i, row_data in enumerate(interface_data, 1):
        for j, val in enumerate(row_data):
            interface_table.rows[i].cells[j].text = val
    
    doc.add_heading('2.4 样品软件信息', level=2)
    sw_table = doc.add_table(rows=8, cols=3)
    sw_table.style = 'Table Grid'
    sw_headers = ['软件组件', '版本号', '说明']
    for i, h in enumerate(sw_headers):
        cell = sw_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    sw_data = [
        ('应用软件', '[版本]', '主控制程序'),
        ('底层驱动', '[版本]', '底层驱动程序'),
        ('RTOS', '[版本]', '实时操作系统'),
        ('通信协议栈', '[版本]', 'CAN/UDS协议栈'),
        ('安全组件', '[版本]', '加密/安全模块'),
        ('Bootloader', '[版本]', '引导加载程序'),
        ('HSM固件', '[版本]', '安全模块固件')
    ]
    for i, row_data in enumerate(sw_data, 1):
        for j, val in enumerate(row_data):
            sw_table.rows[i].cells[j].text = val
    
    doc.add_page_break()
    
    # 3. 测试人员清单
    doc.add_heading('3. 测试人员清单', level=1)
    staff_table = doc.add_table(rows=6, cols=5)
    staff_table.style = 'Table Grid'
    staff_headers = ['序号', '姓名', '角色', '职责', '联系方式']
    for i, h in enumerate(staff_headers):
        cell = staff_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    staff_data = [
        ('1', '[姓名]', '测试负责人', '测试策划、进度管控、报告审核', '[联系方式]'),
        ('2', '[姓名]', '测试工程师', '测试执行、记录、问题跟踪', '[联系方式]'),
        ('3', '[姓名]', '测试工程师', '测试执行、记录、问题跟踪', '[联系方式]'),
        ('4', '[姓名]', '安全专家', '技术指导、复杂问题分析', '[联系方式]'),
        ('5', '[姓名]', '质量工程师', '过程监督、合规性审查', '[联系方式]')
    ]
    for i, row_data in enumerate(staff_data, 1):
        for j, val in enumerate(row_data):
            staff_table.rows[i].cells[j].text = val
    
    doc.add_page_break()
    
    # 4. 测试设备清单
    doc.add_heading('4. 测试设备清单', level=1)
    
    doc.add_heading('4.1 硬件设备', level=2)
    hw_table = doc.add_table(rows=13, cols=6)
    hw_table.style = 'Table Grid'
    hw_headers = ['序号', '设备名称', '型号', '设备编号', '校准有效期', '用途']
    for i, h in enumerate(hw_headers):
        cell = hw_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    hw_data = [
        ('1', 'CAN总线分析仪', 'Vector VN1630', '[编号]', '[日期]', 'CAN/CAN-FD通信测试'),
        ('2', '以太网测试仪', 'Spirent TestCenter', '[编号]', '[日期]', '以太网通信测试'),
        ('3', '示波器', 'Tektronix MSO58', '[编号]', '[日期]', '信号分析'),
        ('4', '逻辑分析仪', 'Saleae Logic Pro 16', '[编号]', '[日期]', '数字信号分析'),
        ('5', '电源供应器', 'Keysight N6705C', '[编号]', '[日期]', '供电测试'),
        ('6', '程控电源', 'EA PSB 10000', '[编号]', '[日期]', '电源波动测试'),
        ('7', '万用表', 'Fluke 87V', '[编号]', '[日期]', '电压测量'),
        ('8', '焊接工具', '[型号]', '[编号]', '-', '硬件调试'),
        ('9', '显微镜', '[型号]', '[编号]', '[日期]', '芯片分析'),
        ('10', '信号发生器', 'Keysight 33622A', '[编号]', '[日期]', '信号模拟'),
        ('11', '故障注入设备', '[型号]', '[编号]', '[日期]', '故障注入测试'),
        ('12', '侧信道分析设备', 'ChipWhisperer', '[编号]', '[日期]', '侧信道攻击测试')
    ]
    for i, row_data in enumerate(hw_data, 1):
        for j, val in enumerate(row_data):
            hw_table.rows[i].cells[j].text = val
    
    doc.add_heading('4.2 软件工具', level=2)
    sw_tool_table = doc.add_table(rows=13, cols=4)
    sw_tool_table.style = 'Table Grid'
    sw_tool_headers = ['序号', '工具名称', '版本', '用途']
    for i, h in enumerate(sw_tool_headers):
        cell = sw_tool_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    sw_tool_data = [
        ('1', 'CANoe', '16.0 SP3', 'CAN总线仿真测试'),
        ('2', 'CANalyzer', '16.0 SP3', 'CAN报文分析'),
        ('3', 'PCAN-Explorer', '6.x', 'CAN总线监控'),
        ('4', 'Wireshark', '4.x', '网络协议分析'),
        ('5', 'JTAG调试器', '[版本]', '固件调试'),
        ('6', 'IDA Pro', '7.x', '固件逆向分析'),
        ('7', 'Ghidra', '10.x', '固件逆向分析'),
        ('8', 'Binwalk', '3.x', '固件提取分析'),
        ('9', 'OpenSSL', '3.x', '加密算法验证'),
        ('10', 'Python', '3.x', '自动化测试脚本'),
        ('11', 'Vector vTESTstudio', '[版本]', '自动化测试开发'),
        ('12', '漏洞扫描工具', '[名称/版本]', '安全漏洞扫描')
    ]
    for i, row_data in enumerate(sw_tool_data, 1):
        for j, val in enumerate(row_data):
            sw_tool_table.rows[i].cells[j].text = val
    
    doc.add_page_break()
    
    # 5. 测试结论
    doc.add_heading('5. 测试结论', level=1)
    
    doc.add_heading('5.1 测试结论总览', level=2)
    conclusion_table = doc.add_table(rows=8, cols=6)
    conclusion_table.style = 'Table Grid'
    conclusion_headers = ['测试类别', '测试项总数', '通过', '不通过', '不适用', '通过率']
    for i, h in enumerate(conclusion_headers):
        cell = conclusion_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    conclusion_data = [
        ('车载软件升级安全', '[X]', '[X]', '[X]', '[X]', '[XX]%'),
        ('车载通信安全', '[X]', '[X]', '[X]', '[X]', '[XX]%'),
        ('外部接口安全', '[X]', '[X]', '[X]', '[X]', '[XX]%'),
        ('数据安全', '[X]', '[X]', '[X]', '[X]', '[XX]%'),
        ('硬件安全', '[X]', '[X]', '[X]', '[X]', '[XX]%'),
        ('软件安全', '[X]', '[X]', '[X]', '[X]', '[XX]%'),
        ('合计', '[X]', '[X]', '[X]', '[X]', '[XX]%')
    ]
    for i, row_data in enumerate(conclusion_data, 1):
        for j, val in enumerate(row_data):
            conclusion_table.rows[i].cells[j].text = val
    
    doc.add_heading('5.2 总体结论', level=2)
    doc.add_paragraph('□ 通过 - 样品满足 GB 44495-2024 标准要求')
    doc.add_paragraph('□ 有条件通过 - 样品基本满足要求，存在遗留问题需整改')
    doc.add_paragraph('□ 不通过 - 样品不满足标准要求，需重大整改后重新测试')
    doc.add_paragraph()
    doc.add_paragraph('结论说明：')
    doc.add_paragraph('[详细说明测试结论，包括主要符合项、主要问题项、风险评估等]')
    
    doc.add_heading('5.3 合规性声明', level=2)
    doc.add_paragraph('本测试报告基于 GB 44495-2024《汽车整车信息安全技术要求》标准要求，对电驱MCU进行了全面的信息安全测试。测试结果表明：')
    doc.add_paragraph('1. [合规项1]')
    doc.add_paragraph('2. [合规项2]')
    doc.add_paragraph('3. [待改进项]')
    
    doc.add_page_break()
    
    # 6. 问题清单
    doc.add_heading('6. 问题清单', level=1)
    
    doc.add_heading('6.1 已解决问题清单', level=2)
    resolved_table = doc.add_table(rows=3, cols=8)
    resolved_table.style = 'Table Grid'
    resolved_headers = ['序号', '问题编号', '问题描述', '严重等级', '发现日期', '解决日期', '解决方案', '验证结果']
    for i, h in enumerate(resolved_headers):
        cell = resolved_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    resolved_table.rows[1].cells[0].text = '1'
    resolved_table.rows[1].cells[1].text = 'BUG-001'
    resolved_table.rows[1].cells[2].text = '[问题描述]'
    resolved_table.rows[1].cells[3].text = '高/中/低'
    resolved_table.rows[1].cells[4].text = '[日期]'
    resolved_table.rows[1].cells[5].text = '[日期]'
    resolved_table.rows[1].cells[6].text = '[解决方案]'
    resolved_table.rows[1].cells[7].text = '已验证通过'
    
    doc.add_heading('6.2 遗留问题清单', level=2)
    pending_table = doc.add_table(rows=3, cols=9)
    pending_table.style = 'Table Grid'
    pending_headers = ['序号', '问题编号', '问题描述', '严重等级', '发现日期', '影响范围', '风险评估', '整改计划', '责任人']
    for i, h in enumerate(pending_headers):
        cell = pending_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    pending_table.rows[1].cells[0].text = '1'
    pending_table.rows[1].cells[1].text = 'BUG-XXX'
    pending_table.rows[1].cells[2].text = '[问题描述]'
    pending_table.rows[1].cells[3].text = '高/中/低'
    pending_table.rows[1].cells[4].text = '[日期]'
    pending_table.rows[1].cells[5].text = '[影响范围]'
    pending_table.rows[1].cells[6].text = '[风险说明]'
    pending_table.rows[1].cells[7].text = '[计划完成日期]'
    pending_table.rows[1].cells[8].text = '[责任人]'
    
    doc.add_heading('6.3 问题严重等级定义', level=2)
    severity_table = doc.add_table(rows=4, cols=3)
    severity_table.style = 'Table Grid'
    severity_headers = ['等级', '定义', '处理要求']
    for i, h in enumerate(severity_headers):
        cell = severity_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    severity_data = [
        ('高', '存在重大安全隐患，可能导致车辆被非法控制、数据泄露等严重后果', '必须立即整改，整改后重新测试'),
        ('中', '存在一般安全隐患，可能影响信息安全防护能力', '应在产品SOP前整改完成'),
        ('低', '存在轻微安全隐患或不符合最佳实践', '建议整改，可纳入后续版本优化')
    ]
    for i, row_data in enumerate(severity_data, 1):
        for j, val in enumerate(row_data):
            severity_table.rows[i].cells[j].text = val
    
    doc.add_page_break()
    
    # 7. 测试详情
    doc.add_heading('7. 测试详情', level=1)
    
    # 7.1 车载软件升级安全测试
    doc.add_heading('7.1 车载软件升级安全测试', level=2)
    
    # 测试项模板函数
    def add_test_item(doc, item_id, item_name, clause, purpose, method_list, step_data, test_data=None, conclusion='□通过 □不通过 □不适用'):
        doc.add_heading(f'测试项编号：{item_id}', level=3)
        doc.add_paragraph(f'测试项名称：{item_name}')
        doc.add_paragraph(f'标准条款：{clause}')
        doc.add_paragraph()
        
        doc.add_paragraph('测试目的：')
        doc.add_paragraph(purpose)
        doc.add_paragraph()
        
        doc.add_paragraph('测试方法：')
        for i, method in enumerate(method_list, 1):
            doc.add_paragraph(f'{i}. {method}')
        doc.add_paragraph()
        
        doc.add_paragraph('测试步骤：')
        step_table = doc.add_table(rows=len(step_data)+1, cols=5)
        step_table.style = 'Table Grid'
        step_headers = ['步骤', '操作描述', '预期结果', '实际结果', '结论']
        for i, h in enumerate(step_headers):
            cell = step_table.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'D9E2F3')
        for i, row_data in enumerate(step_data, 1):
            for j, val in enumerate(row_data):
                step_table.rows[i].cells[j].text = val
        
        if test_data:
            doc.add_paragraph()
            doc.add_paragraph('测试数据：')
            for k, v in test_data.items():
                doc.add_paragraph(f'- {k}：{v}')
        
        doc.add_paragraph()
        doc.add_paragraph(f'测试结论：{conclusion}')
        doc.add_paragraph('问题描述：[如有问题，详细描述]')
        doc.add_paragraph()
    
    # SEC-OTA-001
    add_test_item(doc,
        'SEC-OTA-001',
        '升级包签名验证',
        'GB 44495-2024 第X章X节',
        '验证MCU在接收升级包时是否能够正确验证升级包的数字签名，防止被篡改或伪造的升级包刷写到设备中。',
        [
            '构造合法签名的升级包，验证是否能够正常刷写',
            '构造签名被篡改的升级包，验证是否能够被拒绝',
            '构造无签名的升级包，验证是否能够被拒绝',
            '构造签名证书过期的升级包，验证是否能够被拒绝',
            '构造签名证书吊销的升级包，验证是否能够被拒绝'
        ],
        [
            ('1', '使用合法签名升级包进行刷写', '刷写成功', '', '□通过 □不通过'),
            ('2', '修改升级包内容后刷写', '刷写被拒绝，记录安全日志', '', '□通过 □不通过'),
            ('3', '删除升级包签名后刷写', '刷写被拒绝，记录安全日志', '', '□通过 □不通过'),
            ('4', '使用过期证书签名的升级包刷写', '刷写被拒绝，记录安全日志', '', '□通过 □不通过'),
            ('5', '使用吊销证书签名的升级包刷写', '刷写被拒绝，记录安全日志', '', '□通过 □不通过')
        ],
        {'签名算法': '如RSA-2048/ECDSA-P256', '证书链': '描述证书链结构', '测试样本数量': '[X]个'}
    )
    
    # SEC-OTA-002
    add_test_item(doc,
        'SEC-OTA-002',
        '升级包完整性验证',
        'GB 44495-2024 第X章X节',
        '验证MCU在接收升级包时是否能够验证升级包的完整性，防止升级包在传输过程中被篡改或损坏。',
        [
            '发送完整升级包，验证刷写成功',
            '发送部分损坏的升级包，验证是否被拒绝',
            '发送被篡改的升级包，验证是否被拒绝',
            '模拟传输中断场景，验证恢复机制'
        ],
        [
            ('1', '发送完整升级包', '刷写成功', '', '□通过 □不通过'),
            ('2', '修改升级包1字节数据', '刷写被拒绝，完整性校验失败', '', '□通过 □不通过'),
            ('3', '截断升级包末尾数据', '刷写被拒绝，完整性校验失败', '', '□通过 □不通过'),
            ('4', '传输过程中断电重启', '能够恢复或重新下载', '', '□通过 □不通过')
        ]
    )
    
    # SEC-OTA-003
    add_test_item(doc,
        'SEC-OTA-003',
        '升级包来源认证',
        'GB 44495-2024 第X章X节',
        '验证MCU是否能够验证升级包来源的合法性，防止来自非授权源的升级包刷写。',
        [
            '使用官方证书签名的升级包刷写',
            '使用非官方证书签名的升级包刷写',
            '使用自签名证书签名的升级包刷写'
        ],
        [
            ('1', '官方证书签名升级包', '刷写成功', '', '□通过 □不通过'),
            ('2', '非官方证书签名升级包', '刷写被拒绝', '', '□通过 □不通过'),
            ('3', '自签名证书签名升级包', '刷写被拒绝', '', '□通过 □不通过')
        ]
    )
    
    # SEC-OTA-004
    add_test_item(doc,
        'SEC-OTA-004',
        '升级过程安全状态',
        'GB 44495-2024 第X章X节',
        '验证MCU在升级过程中是否能够保持安全状态，防止升级失败导致设备不可用或进入不安全状态。',
        [
            '正常升级流程测试',
            '升级过程中断电测试',
            '升级过程中通信中断测试',
            '升级失败回滚测试'
        ],
        [
            ('1', '正常升级流程', '升级成功，功能正常', '', '□通过 □不通过'),
            ('2', '升级过程中断电', '重启后可回滚或恢复', '', '□通过 □不通过'),
            ('3', '升级过程通信中断', '能够恢复或回滚', '', '□通过 □不通过'),
            ('4', '强制中断升级', '能够回滚到原版本', '', '□通过 □不通过')
        ]
    )
    
    # SEC-OTA-005
    add_test_item(doc,
        'SEC-OTA-005',
        '升级权限控制',
        'GB 44495-2024 第X章X节',
        '验证MCU是否对升级操作进行权限控制，防止未授权的升级操作。',
        [
            '验证是否需要安全访问（Security Access）才能升级',
            '验证安全访问种子-密钥机制',
            '验证升级会话管理'
        ],
        [
            ('1', '未解锁安全访问直接升级', '升级被拒绝', '', '□通过 □不通过'),
            ('2', '错误解锁密钥升级', '升级被拒绝', '', '□通过 □不通过'),
            ('3', '安全访问超时后升级', '升级被拒绝', '', '□通过 □不通过'),
            ('4', '正确解锁后升级', '升级成功', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 7.1.2 Bootloader安全测试
    doc.add_heading('7.1.2 Bootloader安全测试', level=3)
    
    add_test_item(doc,
        'SEC-OTA-006',
        'Bootloader完整性保护',
        'GB 44495-2024 第X章X节',
        '验证Bootloader自身是否具有完整性保护机制，防止被篡改。',
        [
            '检查Bootloader是否有签名/校验机制',
            '尝试修改Bootloader并验证是否能够检测',
            '验证Bootloader启动时的自校验'
        ],
        [
            ('1', '检查Bootloader保护机制', '存在完整性保护机制', '', '□通过 □不通过'),
            ('2', '修改Bootloader后启动', '启动失败或报警', '', '□通过 □不通过'),
            ('3', '正常Bootloader启动', '启动成功', '', '□通过 □不通过')
        ]
    )
    
    add_test_item(doc,
        'SEC-OTA-007',
        'Bootloader安全启动',
        'GB 44495-2024 第X章X节',
        '验证MCU是否支持安全启动（Secure Boot），确保启动链完整性。',
        [
            '验证是否支持安全启动',
            '尝试加载未签名应用程序',
            '尝试加载被篡改应用程序'
        ],
        [
            ('1', '检查安全启动配置', '已启用安全启动', '', '□通过 □不通过'),
            ('2', '加载未签名应用', '启动被拒绝', '', '□通过 □不通过'),
            ('3', '加载篡改应用', '启动被拒绝', '', '□通过 □不通过'),
            ('4', '加载合法应用', '启动成功', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 7.2 车载通信安全测试
    doc.add_heading('7.2 车载通信安全测试', level=2)
    doc.add_heading('7.2.1 CAN总线通信安全测试', level=3)
    
    add_test_item(doc,
        'SEC-COM-001',
        'CAN报文认证测试',
        'GB 44495-2024 第X章X节',
        '验证MCU发送的CAN报文是否具有认证机制，防止报文被伪造或重放。',
        [
            '分析CAN报文格式，检查是否有认证字段',
            '构造伪造CAN报文发送给MCU',
            '重放历史CAN报文',
            '检查是否支持SecOC（安全车载通信）'
        ],
        [
            ('1', '分析CAN报文格式', '存在认证字段/MAC', '', '□通过 □不通过'),
            ('2', '发送伪造报文', 'MCU正确拒绝或忽略', '', '□通过 □不通过'),
            ('3', '重放历史报文', 'MCU检测到重放攻击', '', '□通过 □不通过'),
            ('4', '检查SecOC实现', '已实现SecOC机制', '', '□通过 □不通过')
        ],
        {'CAN ID列表': '列出关键CAN ID', '认证算法': '如AES-CMAC', 'MAC长度': '[X]bit'}
    )
    
    add_test_item(doc,
        'SEC-COM-002',
        'CAN报文完整性保护',
        'GB 44495-2024 第X章X节',
        '验证CAN报文是否具有完整性保护机制，防止报文在传输过程中被篡改。',
        [
            '监听正常CAN报文',
            '修改报文数据字段后发送',
            '验证MCU是否能够检测篡改'
        ],
        [
            ('1', '监听正常报文', '成功捕获报文', '', '□通过 □不通过'),
            ('2', '修改数据字段发送', 'MCU检测到篡改', '', '□通过 □不通过'),
            ('3', '修改认证字段发送', 'MCU检测到篡改', '', '□通过 □不通过')
        ]
    )
    
    add_test_item(doc,
        'SEC-COM-003',
        'CAN总线DoS攻击防护',
        'GB 44495-2024 第X章X节',
        '验证MCU是否能够抵御CAN总线DoS攻击，保持基本功能可用。',
        [
            '发送大量高优先级报文占用总线',
            '发送错误帧导致总线错误',
            '模拟总线风暴攻击',
            '验证MCU的防护机制和恢复能力'
        ],
        [
            ('1', '总线高负载测试', 'MCU功能正常', '', '□通过 □不通过'),
            ('2', '错误帧注入测试', 'MCU能够恢复', '', '□通过 □不通过'),
            ('3', '总线风暴攻击', 'MCU能够检测并保护', '', '□通过 □不通过'),
            ('4', '攻击停止后恢复', 'MCU能够自动恢复', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 7.2.2 诊断通信安全测试
    doc.add_heading('7.2.2 诊断通信安全测试', level=3)
    
    add_test_item(doc,
        'SEC-COM-005',
        'UDS安全访问测试',
        'GB 44495-2024 第X章X节',
        '验证MCU的UDS诊断服务是否具有安全访问控制机制。',
        [
            '测试安全访问服务（0x27服务）',
            '验证种子-密钥算法',
            '测试安全等级划分',
            '测试会话超时机制'
        ],
        [
            ('1', '请求安全访问种子', '返回随机种子', '', '□通过 □不通过'),
            ('2', '发送错误密钥', '访问被拒绝', '', '□通过 □不通过'),
            ('3', '发送正确密钥', '访问成功', '', '□通过 □不通过'),
            ('4', '安全会话超时测试', '超时后需重新解锁', '', '□通过 □不通过'),
            ('5', '连续错误尝试测试', '存在防暴力破解机制', '', '□通过 □不通过')
        ],
        {'安全等级': '列出各安全等级', '种子长度': '[X]字节', '密钥长度': '[X]字节', '最大错误尝试次数': '[X]次'}
    )
    
    add_test_item(doc,
        'SEC-COM-006',
        '诊断会话管理测试',
        'GB 44495-2024 第X章X节',
        '验证MCU的诊断会话管理是否安全。',
        [
            '测试会话切换',
            '测试会话超时',
            '测试会话权限隔离'
        ],
        [
            ('1', '默认会话权限测试', '仅允许基本诊断', '', '□通过 □不通过'),
            ('2', '编程会话权限测试', '需安全访问', '', '□通过 □不通过'),
            ('3', '扩展会话权限测试', '需安全访问', '', '□通过 □不通过'),
            ('4', '会话超时测试', '超时返回默认会话', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 7.3 外部接口安全测试
    doc.add_heading('7.3 外部接口安全测试', level=2)
    doc.add_heading('7.3.1 调试接口安全测试', level=3)
    
    add_test_item(doc,
        'SEC-INT-001',
        'JTAG/SWD接口保护测试',
        'GB 44495-2024 第X章X节',
        '验证MCU的JTAG/SWD调试接口是否受到保护，防止通过调试接口获取固件或篡改程序。',
        [
            '检查JTAG/SWD接口物理存在性',
            '尝试连接调试器',
            '尝试读取Flash内容',
            '尝试写入/擦除Flash',
            '检查调试接口禁用机制'
        ],
        [
            ('1', '检查调试接口状态', '已禁用或受保护', '', '□通过 □不通过'),
            ('2', '尝试连接调试器', '连接被拒绝或需认证', '', '□通过 □不通过'),
            ('3', '尝试读取Flash', '读取被拒绝或返回无效数据', '', '□通过 □不通过'),
            ('4', '尝试写入Flash', '写入被拒绝', '', '□通过 □不通过'),
            ('5', '检查JTAG熔丝位', '已熔断或已配置保护', '', '□通过 □不通过')
        ]
    )
    
    add_test_item(doc,
        'SEC-INT-002',
        '串口调试接口保护测试',
        'GB 44495-2024 第X章X节',
        '验证MCU的串口调试接口是否受到保护。',
        [
            '检查串口调试接口',
            '尝试访问调试Shell',
            '尝试通过串口刷写固件'
        ],
        [
            ('1', '检查串口输出', '无敏感信息输出', '', '□通过 □不通过'),
            ('2', '尝试访问调试Shell', '无调试Shell或需认证', '', '□通过 □不通过'),
            ('3', '尝试串口刷写', '需安全访问验证', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 7.4 数据安全测试
    doc.add_heading('7.4 数据安全测试', level=2)
    doc.add_heading('7.4.1 敏感数据存储安全测试', level=3)
    
    add_test_item(doc,
        'SEC-DATA-001',
        '密钥存储安全测试',
        'GB 44495-2024 第X章X节',
        '验证MCU中密钥的存储是否安全，防止密钥被提取。',
        [
            '分析密钥存储位置',
            '尝试通过调试接口读取密钥',
            '尝试通过诊断接口读取密钥',
            '检查HSM/SE使用情况'
        ],
        [
            ('1', '检查密钥存储位置', '存储在HSM/SE中', '', '□通过 □不通过'),
            ('2', '调试接口读取测试', '无法读取密钥', '', '□通过 □不通过'),
            ('3', '诊断接口读取测试', '无法读取密钥', '', '□通过 □不通过'),
            ('4', '密钥导出测试', '密钥不可导出', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 7.5 硬件安全测试
    doc.add_heading('7.5 硬件安全测试', level=2)
    doc.add_heading('7.5.1 芯片安全测试', level=3)
    
    add_test_item(doc,
        'SEC-HW-001',
        '芯片安全功能测试',
        'GB 44495-2024 第X章X节',
        '验证MCU芯片是否具备安全功能。',
        [
            '检查芯片安全功能列表',
            '测试HSM功能',
            '测试安全启动功能',
            '测试调试保护功能'
        ],
        [
            ('1', '检查HSM模块', '存在HSM模块', '', '□通过 □不通过'),
            ('2', '测试HSM加密功能', '加解密功能正常', '', '□通过 □不通过'),
            ('3', '测试安全启动', '安全启动已启用', '', '□通过 □不通过'),
            ('4', '测试调试保护', '调试接口已保护', '', '□通过 □不通过')
        ]
    )
    
    add_test_item(doc,
        'SEC-HW-002',
        '侧信道攻击防护测试',
        'GB 44495-2024 第X章X节',
        '验证MCU是否具有侧信道攻击防护能力。',
        [
            '功耗分析攻击测试',
            '电磁辐射分析攻击测试',
            '时序分析攻击测试'
        ],
        [
            ('1', '功耗分析测试', '无法提取密钥', '', '□通过 □不通过'),
            ('2', '电磁分析测试', '无法提取密钥', '', '□通过 □不通过'),
            ('3', '时序分析测试', '无法提取密钥', '', '□通过 □不通过')
        ]
    )
    
    add_test_item(doc,
        'SEC-HW-003',
        '故障注入攻击防护测试',
        'GB 44495-2024 第X章X节',
        '验证MCU是否具有故障注入攻击防护能力。',
        [
            '电压故障注入测试',
            '时钟故障注入测试',
            '电磁故障注入测试'
        ],
        [
            ('1', '电压故障注入', '检测到攻击并保护', '', '□通过 □不通过'),
            ('2', '时钟故障注入', '检测到攻击并保护', '', '□通过 □不通过'),
            ('3', '电磁故障注入', '检测到攻击并保护', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 7.6 软件安全测试
    doc.add_heading('7.6 软件安全测试', level=2)
    doc.add_heading('7.6.1 固件安全测试', level=3)
    
    add_test_item(doc,
        'SEC-SW-001',
        '固件逆向分析测试',
        'GB 44495-2024 第X章X节',
        '验证固件是否具有防逆向分析保护。',
        [
            '提取固件（如可提取）',
            '分析固件结构',
            '尝试反编译分析'
        ],
        [
            ('1', '固件提取测试', '无法提取或提取困难', '', '□通过 □不通过'),
            ('2', '固件加密分析', '固件已加密', '', '□通过 □不通过'),
            ('3', '代码混淆分析', '存在代码混淆', '', '□通过 □不通过'),
            ('4', '字符串分析', '敏感字符串已混淆', '', '□通过 □不通过')
        ]
    )
    
    add_test_item(doc,
        'SEC-SW-004',
        '加密算法安全性测试',
        'GB 44495-2024 第X章X节',
        '验证MCU使用的加密算法是否安全。',
        [
            '识别使用的加密算法',
            '验证算法参数',
            '测试算法实现正确性'
        ],
        [
            ('1', '识别加密算法', '使用安全加密算法', '', '□通过 □不通过'),
            ('2', 'AES密钥长度检查', '密钥长度≥128位', '', '□通过 □不通过'),
            ('3', 'RSA密钥长度检查', '密钥长度≥2048位', '', '□通过 □不通过'),
            ('4', '哈希算法检查', '使用SHA-256及以上', '', '□通过 □不通过')
        ],
        {'对称加密算法': '如AES-128-GCM', '非对称加密算法': '如RSA-2048/ECDSA-P256', '哈希算法': '如SHA-256', 'MAC算法': '如HMAC-SHA256/AES-CMAC'}
    )
    
    add_test_item(doc,
        'SEC-SW-006',
        '安全日志记录测试',
        'GB 44495-2024 第X章X节',
        '验证MCU是否正确记录安全事件日志。',
        [
            '触发各类安全事件',
            '检查日志记录',
            '验证日志完整性'
        ],
        [
            ('1', '触发安全访问失败', '记录安全日志', '', '□通过 □不通过'),
            ('2', '触发异常刷写尝试', '记录安全日志', '', '□通过 □不通过'),
            ('3', '触发通信异常', '记录安全日志', '', '□通过 □不通过'),
            ('4', '日志完整性验证', '日志不可篡改', '', '□通过 □不通过')
        ]
    )
    
    doc.add_page_break()
    
    # 附录
    doc.add_heading('附录', level=1)
    
    doc.add_heading('附录A：测试环境配置记录', level=2)
    doc.add_paragraph('[详细记录测试环境配置信息]')
    
    doc.add_heading('附录B：测试报文记录', level=2)
    doc.add_paragraph('[记录关键测试报文数据]')
    
    doc.add_heading('附录C：测试日志', level=2)
    doc.add_paragraph('[附上测试过程中的日志文件]')
    
    doc.add_heading('附录D：测试照片', level=2)
    doc.add_paragraph('[附上测试过程中的关键照片]')
    
    doc.add_heading('附录E：参考文档', level=2)
    ref_items = [
        '1. GB 44495-2024《汽车整车信息安全技术要求》',
        '2. GB/T 40857-2021《汽车信息安全通用技术要求》',
        '3. ISO/SAE 21434:2021 Road vehicles – Cybersecurity engineering',
        '4. UN R155 Uniform provisions concerning the approval of vehicles with regards to cyber security'
    ]
    for item in ref_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # 签署页
    doc.add_heading('签署页', level=1)
    sign_table = doc.add_table(rows=4, cols=4)
    sign_table.style = 'Table Grid'
    sign_headers = ['角色', '姓名', '签名', '日期']
    for i, h in enumerate(sign_headers):
        cell = sign_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'D9E2F3')
    sign_data = [
        ('编制', '', '', ''),
        ('审核', '', '', ''),
        ('批准', '', '', '')
    ]
    for i, row_data in enumerate(sign_data, 1):
        for j, val in enumerate(row_data):
            sign_table.rows[i].cells[j].text = val
    
    doc.add_paragraph()
    doc.add_paragraph('报告结束', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    doc.save(r'D:\AI\GB44495-2024_MCU信息安全测试报告模板.docx')
    print('Word文档已生成：D:\\AI\\GB44495-2024_MCU信息安全测试报告模板.docx')

if __name__ == '__main__':
    create_report()
