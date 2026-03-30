# -*- coding: utf-8 -*-
"""
PDF提取工具
功能：提取PDF中的文本、图片、表格，输出到Word文档
"""

import os
import sys
from pathlib import Path

# 自动安装依赖
def check_and_install_packages():
    packages = {
        'pdfplumber': 'pdfplumber',
        'python-docx': 'docx',
        'Pillow': 'PIL'
    }
    
    missing = []
    for pip_name, import_name in packages.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pip_name)
    
    if missing:
        print(f"正在安装依赖包: {', '.join(missing)}")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q'] + missing)
        print("安装完成！")

check_and_install_packages()

import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_chinese_font(doc):
    """设置中文字体"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def extract_images_from_page(page, page_num, img_dir):
    """从PDF页面提取图片"""
    images = []
    try:
        for i, img in enumerate(page.images):
            try:
                if 'stream' in img:
                    stream = img['stream']
                    if hasattr(stream, 'get_data'):
                        img_data = stream.get_data()
                        if img_data:
                            img_path = img_dir / f'page{page_num}_img{i}.png'
                            with open(img_path, 'wb') as f:
                                f.write(img_data)
                            images.append(str(img_path))
            except Exception:
                continue
    except Exception:
        pass
    return images


def pdf_to_word(pdf_path, output_path=None, extract_images=True):
    """
    PDF转Word
    
    参数:
        pdf_path: PDF文件路径
        output_path: 输出路径（可选）
        extract_images: 是否提取图片
    """
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    if output_path is None:
        output_path = pdf_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    output_dir = output_path.parent
    img_dir = output_dir / 'extracted_images'
    if extract_images:
        img_dir.mkdir(exist_ok=True)
    
    print(f"正在处理: {pdf_path.name}")
    
    doc = Document()
    set_chinese_font(doc)
    
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        print(f"共 {total_pages} 页")
        
        for page_num, page in enumerate(pdf.pages, 1):
            print(f"  处理第 {page_num}/{total_pages} 页...", end='\r')
            
            # 页码标记
            if total_pages > 1:
                p = doc.add_paragraph()
                p.add_run(f"—— 第 {page_num} 页 ——").bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 提取文本
            text = page.extract_text()
            if text:
                for para_text in text.split('\n'):
                    para_text = para_text.strip()
                    if para_text:
                        doc.add_paragraph(para_text)
            
            # 提取表格
            tables = page.extract_tables()
            if tables:
                for table_data in tables:
                    if table_data and len(table_data) > 0:
                        rows = len(table_data)
                        cols = max(len(row) for row in table_data)
                        if rows > 0 and cols > 0:
                            table = doc.add_table(rows=rows, cols=cols)
                            table.style = 'Table Grid'
                            for i, row_data in enumerate(table_data):
                                for j, cell_data in enumerate(row_data):
                                    if j < cols and cell_data:
                                        table.rows[i].cells[j].text = str(cell_data).strip()
                            doc.add_paragraph()
            
            # 提取图片
            if extract_images:
                images = extract_images_from_page(page, page_num, img_dir)
                for img_path in images:
                    try:
                        doc.add_picture(img_path, width=Inches(4))
                        doc.add_paragraph()
                    except Exception:
                        pass
            
            if page_num < total_pages:
                doc.add_paragraph()
        
        print(f"\n处理完成！")
    
    doc.save(output_path)
    print(f"Word文档已保存: {output_path}")
    return str(output_path)


def extract_text_only(pdf_path, output_path=None):
    """仅提取文本"""
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    if output_path is None:
        output_path = pdf_path.with_suffix('.txt')
    else:
        output_path = Path(output_path)
    
    print(f"正在提取文本: {pdf_path.name}")
    
    all_text = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                all_text.append(f"=== 第 {page_num} 页 ===\n{text}")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(all_text))
    
    print(f"文本已保存: {output_path}")
    return str(output_path)


def extract_images_only(pdf_path, output_dir=None):
    """仅提取图片"""
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    if output_dir is None:
        output_dir = pdf_path.parent / f"{pdf_path.stem}_images"
    else:
        output_dir = Path(output_dir)
    
    output_dir.mkdir(exist_ok=True)
    
    print(f"正在提取图片: {pdf_path.name}")
    
    count = 0
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            try:
                for i, img in enumerate(page.images):
                    try:
                        if 'stream' in img:
                            stream = img['stream']
                            if hasattr(stream, 'get_data'):
                                img_data = stream.get_data()
                                if img_data:
                                    img_path = output_dir / f'page{page_num}_img{i}.png'
                                    with open(img_path, 'wb') as f:
                                        f.write(img_data)
                                    count += 1
                    except Exception:
                        continue
            except Exception:
                continue
    
    print(f"共提取 {count} 张图片到: {output_dir}")
    return str(output_dir)


def extract_tables_only(pdf_path, output_dir=None):
    """仅提取表格"""
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    if output_dir is None:
        output_dir = pdf_path.parent / f"{pdf_path.stem}_tables"
    else:
        output_dir = Path(output_dir)
    
    output_dir.mkdir(exist_ok=True)
    
    print(f"正在提取表格: {pdf_path.name}")
    
    count = 0
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            tables = page.extract_tables()
            if tables:
                for i, table_data in enumerate(tables):
                    if table_data and len(table_data) > 0:
                        import csv
                        csv_path = output_dir / f'page{page_num}_table{i}.csv'
                        with open(csv_path, 'w', encoding='utf-8-sig', newline='') as f:
                            writer = csv.writer(f)
                            writer.writerows(table_data)
                        count += 1
    
    print(f"共提取 {count} 个表格到: {output_dir}")
    return str(output_dir)


if __name__ == '__main__':
    # ==================== 配置区域 ====================
    # PDF文件路径
    PDF_PATH = r"D:\AI\test.pdf"
    
    # 输出路径（None表示与PDF同目录）
    OUTPUT_PATH = None
    
    # 提取模式: "all" / "text" / "images" / "tables"
    MODE = "all"
    
    # 是否提取图片（仅MODE="all"时有效）
    EXTRACT_IMAGES = True
    # ================================================
    
    if MODE == "all":
        pdf_to_word(PDF_PATH, OUTPUT_PATH, EXTRACT_IMAGES)
    elif MODE == "text":
        extract_text_only(PDF_PATH, OUTPUT_PATH)
    elif MODE == "images":
        extract_images_only(PDF_PATH, OUTPUT_PATH)
    elif MODE == "tables":
        extract_tables_only(PDF_PATH, OUTPUT_PATH)
    else:
        print(f"未知模式: {MODE}")
