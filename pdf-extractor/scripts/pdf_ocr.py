# -*- coding: utf-8 -*-
"""
PDF OCR提取工具
功能：对扫描版PDF进行OCR识别，提取文字内容
适用场景：纯图片扫描的PDF，无法直接提取文字
"""

import os
import sys
from pathlib import Path

# 自动安装依赖
def check_and_install_packages():
    packages = {
        'pdf2image': 'pdf2image',
        'paddleocr': 'paddleocr',
        'paddlepaddle': 'paddlepaddle',
        'Pillow': 'PIL'
    }
    
    missing = []
    for pip_name, import_name in packages.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pip_name)
    
    if missing:
        print(f"正在安装依赖包: {', '.join(missing)}")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q'] + missing)
        print("安装完成！")

check_and_install_packages()

from pdf2image import convert_from_path
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def set_chinese_font(doc):
    """设置中文字体"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def pdf_ocr_to_text(pdf_path, output_path=None, lang='ch'):
    """
    PDF OCR识别转文本
    
    参数:
        pdf_path: PDF文件路径
        output_path: 输出路径（可选）
        lang: OCR语言，'ch'=中文，'en'=英文
    """
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    if output_path is None:
        output_path = pdf_path.with_suffix('.txt')
    else:
        output_path = Path(output_path)
    
    print(f"正在处理: {pdf_path.name}")
    print("初始化OCR引擎...")
    
    # 初始化PaddleOCR（首次运行会自动下载模型）
    ocr = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)
    
    # PDF转图片
    print("转换PDF为图片...")
    images = convert_from_path(str(pdf_path), dpi=200)
    total_pages = len(images)
    print(f"共 {total_pages} 页")
    
    all_text = []
    
    for page_num, image in enumerate(images, 1):
        print(f"  OCR识别第 {page_num}/{total_pages} 页...", end='\r')
        
        # 保存临时图片
        temp_img_path = Path(f"temp_page_{page_num}.png")
        image.save(temp_img_path)
        
        # OCR识别
        result = ocr.ocr(str(temp_img_path), cls=True)
        
        # 提取文字
        page_text = []
        if result and result[0]:
            for line in result[0]:
                if line and len(line) >= 2:
                    text = line[1][0]  # 获取识别的文字
                    page_text.append(text)
        
        all_text.append(f"=== 第 {page_num} 页 ===\n" + "\n".join(page_text))
        
        # 删除临时图片
        temp_img_path.unlink(missing_ok=True)
    
    print(f"\n识别完成！")
    
    # 保存文本
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(all_text))
    
    print(f"文本已保存: {output_path}")
    return str(output_path)


def pdf_ocr_to_word(pdf_path, output_path=None, lang='ch'):
    """
    PDF OCR识别转Word文档
    
    参数:
        pdf_path: PDF文件路径
        output_path: 输出路径（可选）
        lang: OCR语言，'ch'=中文，'en'=英文
    """
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    if output_path is None:
        output_path = pdf_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    print(f"正在处理: {pdf_path.name}")
    print("初始化OCR引擎...")
    
    # 初始化PaddleOCR
    ocr = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)
    
    # PDF转图片
    print("转换PDF为图片...")
    images = convert_from_path(str(pdf_path), dpi=200)
    total_pages = len(images)
    print(f"共 {total_pages} 页")
    
    # 创建Word文档
    doc = Document()
    set_chinese_font(doc)
    
    for page_num, image in enumerate(images, 1):
        print(f"  OCR识别第 {page_num}/{total_pages} 页...", end='\r')
        
        # 保存临时图片
        temp_img_path = Path(f"temp_page_{page_num}.png")
        image.save(temp_img_path)
        
        # 页码标记
        if total_pages > 1:
            p = doc.add_paragraph()
            p.add_run(f"—— 第 {page_num} 页 ——").bold = True
            p.alignment = 1  # 居中
        
        # OCR识别
        result = ocr.ocr(str(temp_img_path), cls=True)
        
        # 写入Word
        if result and result[0]:
            for line in result[0]:
                if line and len(line) >= 2:
                    text = line[1][0]
                    doc.add_paragraph(text)
        
        doc.add_paragraph()  # 页间空行
        
        # 删除临时图片
        temp_img_path.unlink(missing_ok=True)
    
    print(f"\n识别完成！")
    
    doc.save(output_path)
    print(f"Word文档已保存: {output_path}")
    return str(output_path)


if __name__ == '__main__':
    # ==================== 配置区域 ====================
    # PDF文件路径
    PDF_PATH = r"D:\AI\scan_test.pdf"
    
    # 输出路径（None表示与PDF同目录）
    OUTPUT_PATH = None
    
    # 输出格式: "text" / "word"
    OUTPUT_FORMAT = "word"
    
    # OCR语言: 'ch'=中文, 'en'=英文
    LANG = "ch"
    # ================================================
    
    if OUTPUT_FORMAT == "text":
        pdf_ocr_to_text(PDF_PATH, OUTPUT_PATH, LANG)
    elif OUTPUT_FORMAT == "word":
        pdf_ocr_to_word(PDF_PATH, OUTPUT_PATH, LANG)
    else:
        print(f"未知格式: {OUTPUT_FORMAT}")
