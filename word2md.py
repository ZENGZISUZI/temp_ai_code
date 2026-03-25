# -*- coding: utf-8 -*-
"""
Word to Markdown Converter
Convert .docx files to clean Markdown format

Author: Assistant
"""

import sys
import os
import re
import argparse
from pathlib import Path

# Fix Windows encoding
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Error: python-docx not installed. Run: pip install python-docx")


class Word2Markdown:
    """Convert Word document to Markdown"""
    
    def __init__(self, extract_images: bool = False, image_dir: str = None):
        self.extract_images = extract_images
        self.image_dir = image_dir
        self.image_count = 0
        self.list_counters = {}  # For nested lists
    
    def convert(self, docx_path: str, output_path: str = None) -> str:
        """
        Convert Word document to Markdown
        
        Args:
            docx_path: Path to .docx file
            output_path: Output .md file path (optional)
        
        Returns:
            Markdown content as string
        """
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")
        
        # Setup image directory
        if self.extract_images and not self.image_dir:
            self.image_dir = Path(docx_path).stem + "_images"
        
        if self.extract_images:
            os.makedirs(self.image_dir, exist_ok=True)
        
        # Load document
        doc = Document(docx_path)
        
        # Convert
        markdown_lines = []
        
        # Process each paragraph
        for para in doc.paragraphs:
            md_line = self._convert_paragraph(para)
            if md_line:
                markdown_lines.append(md_line)
        
        # Process tables
        for table in doc.tables:
            md_table = self._convert_table(table)
            if md_table:
                markdown_lines.append(md_table)
                markdown_lines.append("")  # Add spacing
        
        # Process images if enabled
        if self.extract_images:
            self._extract_images(doc)
        
        # Join and clean up
        markdown = "\n".join(markdown_lines)
        markdown = self._clean_markdown(markdown)
        
        # Save to file if output path specified
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown)
            print(f"[OK] Saved to: {output_path}")
        
        return markdown
    
    def _convert_paragraph(self, para) -> str:
        """Convert a paragraph to Markdown"""
        if not para.text.strip():
            return ""
        
        style_name = para.style.name.lower() if para.style else ""
        text = self._process_runs(para)
        
        # Headings
        if 'heading 1' in style_name or 'title' in style_name:
            return f"# {text}"
        elif 'heading 2' in style_name:
            return f"## {text}"
        elif 'heading 3' in style_name:
            return f"### {text}"
        elif 'heading 4' in style_name:
            return f"#### {text}"
        elif 'heading 5' in style_name:
            return f"##### {text}"
        elif 'heading 6' in style_name:
            return f"###### {text}"
        
        # Lists
        if 'list' in style_name:
            return self._convert_list_item(para, text)
        
        # Check for bullet/number patterns in text
        if text.strip().startswith(('• ', '- ', '* ')):
            return f"- {text.strip()[2:]}"
        
        # Check for numbered pattern
        match = re.match(r'^(\d+)[\.、\)]\s*(.+)$', text.strip())
        if match:
            return f"{match.group(1)}. {match.group(2)}"
        
        # Quote
        if 'quote' in style_name:
            return f"> {text}"
        
        # Code block
        if 'code' in style_name or 'source' in style_name:
            return f"```\n{text}\n```"
        
        # Check alignment for center/right
        alignment = para.alignment
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return f"<div align='center'>{text}</div>"
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return f"<div align='right'>{text}</div>"
        
        # Regular paragraph
        return text
    
    def _process_runs(self, para) -> str:
        """Process runs to preserve formatting"""
        result = []
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # Check formatting
            is_bold = run.bold
            is_italic = run.italic
            is_strike = run.font.strike
            is_code = run.font.name and 'consol' in run.font.name.lower()
            
            # Apply markdown formatting
            if is_code:
                text = f"`{text}`"
            else:
                if is_bold:
                    text = f"**{text}**"
                if is_italic:
                    text = f"*{text}*"
                if is_strike:
                    text = f"~~{text}~~"
            
            result.append(text)
        
        return "".join(result)
    
    def _convert_list_item(self, para, text: str) -> str:
        """Convert list item with proper indentation"""
        # Get list level from paragraph format
        level = 0
        if para.paragraph_format.left_indent:
            # Approximate level from indent (720 twips = 0.5 inch = 1 level)
            level = int(para.paragraph_format.left_indent / 720)
        
        indent = "  " * level
        
        # Check if it's a numbered list
        if para.style and 'number' in para.style.name.lower():
            list_key = f"list_{level}"
            self.list_counters[list_key] = self.list_counters.get(list_key, 0) + 1
            return f"{indent}{self.list_counters[list_key]}. {text}"
        else:
            return f"{indent}- {text}"
    
    def _convert_table(self, table) -> str:
        """Convert table to Markdown"""
        if not table.rows:
            return ""
        
        lines = []
        
        # Get all rows
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # Get cell text, handle merged cells
                cell_text = cell.text.strip().replace('\n', ' ')
                cells.append(cell_text)
            rows.append(cells)
        
        if not rows:
            return ""
        
        # Determine column count
        col_count = max(len(row) for row in rows)
        
        # Build header
        header = rows[0]
        while len(header) < col_count:
            header.append("")
        
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * col_count) + " |")
        
        # Build body
        for row in rows[1:]:
            while len(row) < col_count:
                row.append("")
            lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(lines)
    
    def _extract_images(self, doc):
        """Extract images from document"""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    self.image_count += 1
                    image_data = rel.target_part.blob
                    
                    # Determine extension
                    content_type = rel.target_part.content_type
                    ext = content_type.split('/')[-1]
                    if ext == 'jpeg':
                        ext = 'jpg'
                    
                    # Save image
                    image_name = f"image_{self.image_count}.{ext}"
                    image_path = os.path.join(self.image_dir, image_name)
                    
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                    
                    print(f"[OK] Extracted: {image_path}")
        except Exception as e:
            print(f"[Warning] Could not extract images: {e}")
    
    def _clean_markdown(self, text: str) -> str:
        """Clean up markdown output"""
        # Remove multiple blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Fix spacing around headers
        text = re.sub(r'\n*(#{1,6}\s)', r'\n\n\1', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text


def convert_file(docx_path: str, output_path: str = None, extract_images: bool = False) -> str:
    """
    Convenience function to convert a Word file to Markdown
    
    Args:
        docx_path: Path to .docx file
        output_path: Output .md file path (optional, auto-generated if not specified)
        extract_images: Whether to extract images
    
    Returns:
        Markdown content as string
    """
    converter = Word2Markdown(extract_images=extract_images)
    
    if not output_path:
        output_path = Path(docx_path).stem + ".md"
    
    return converter.convert(docx_path, output_path)


def batch_convert(input_dir: str, output_dir: str = None, extract_images: bool = False):
    """
    Batch convert all Word files in a directory
    
    Args:
        input_dir: Directory containing .docx files
        output_dir: Output directory (optional, same as input if not specified)
        extract_images: Whether to extract images
    """
    if not output_dir:
        output_dir = input_dir
    
    os.makedirs(output_dir, exist_ok=True)
    
    docx_files = list(Path(input_dir).glob("*.docx"))
    
    if not docx_files:
        print(f"No .docx files found in: {input_dir}")
        return
    
    print(f"Found {len(docx_files)} Word files")
    print("=" * 50)
    
    converter = Word2Markdown(extract_images=extract_images)
    
    for docx_file in docx_files:
        try:
            output_path = os.path.join(output_dir, docx_file.stem + ".md")
            print(f"\nConverting: {docx_file.name}")
            converter.convert(str(docx_file), output_path)
        except Exception as e:
            print(f"[Error] Failed to convert {docx_file.name}: {e}")
    
    print("\n" + "=" * 50)
    print("Batch conversion complete!")


def main():
    parser = argparse.ArgumentParser(
        description='Convert Word documents to Markdown',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Single file
  python word2md.py report.docx
  
  # Specify output path
  python word2md.py report.docx -o output.md
  
  # Extract images
  python word2md.py report.docx --images
  
  # Batch convert directory
  python word2md.py ./docs --batch
        """
    )
    
    parser.add_argument('input', help='Input .docx file or directory')
    parser.add_argument('-o', '--output', help='Output .md file or directory')
    parser.add_argument('--images', action='store_true', help='Extract images from document')
    parser.add_argument('--batch', action='store_true', help='Batch convert all .docx files in directory')
    
    args = parser.parse_args()
    
    if not HAS_DOCX:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
        print("Please run the command again.")
        return
    
    if args.batch:
        batch_convert(args.input, args.output, args.images)
    else:
        output = args.output or (Path(args.input).stem + ".md")
        converter = Word2Markdown(extract_images=args.images)
        
        print(f"Converting: {args.input}")
        print("=" * 50)
        
        markdown = converter.convert(args.input, output)
        
        print("\n" + "=" * 50)
        print("Conversion complete!")
        print(f"\nPreview:\n{'-' * 30}")
        print(markdown[:500] + "..." if len(markdown) > 500 else markdown)


if __name__ == '__main__':
    main()
