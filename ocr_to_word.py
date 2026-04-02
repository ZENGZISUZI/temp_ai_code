"""
图片文字识别并生成Word文档
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建Word文档
doc = Document()

# 标题
title = doc.add_heading('Python学习计划', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 第1阶段
doc.add_heading('第1阶段：Python基础（第1-2周）', level=1)
content1 = [
    '1. Python环境搭建与配置',
    '2. 变量、数据类型与运算符',
    '3. 条件语句（if/elif/else）',
    '4. 循环语句（for/while）',
    '5. 列表、元组、字典、集合',
    '6. 函数定义与调用',
    '7. 文件读写操作',
    '',
    '练习题目：',
    '- 计算器程序',
    '- 猜数字游戏',
    '- 学生成绩管理系统'
]
for item in content1:
    doc.add_paragraph(item)

# 第2阶段
doc.add_heading('第2阶段：Python进阶（第3-4周）', level=1)
content2 = [
    '1. 面向对象编程（类与对象）',
    '2. 继承、封装、多态',
    '3. 异常处理',
    '4. 模块与包管理',
    '5. 正则表达式',
    '6. 日期时间处理',
    '7. 多线程与多进程基础',
    '',
    '练习题目：',
    '- 图书管理系统（OOP版）',
    '- 日志分析工具',
    '- 文件批量重命名工具'
]
for item in content2:
    doc.add_paragraph(item)

# 第3阶段
doc.add_heading('第3阶段：Python高级（第5-6周）', level=1)
content3 = [
    '1. 装饰器与生成器',
    '2. 上下文管理器',
    '3. 迭代器与可迭代对象',
    '4. 网络编程基础（socket）',
    '5. 数据库操作（SQLite/MySQL）',
    '6. API调用与JSON处理',
    '7. 单元测试与调试',
    '',
    '练习题目：',
    '- 网页爬虫',
    '- 数据库管理系统',
    '- REST API客户端'
]
for item in content3:
    doc.add_paragraph(item)

# 第4阶段
doc.add_heading('第4阶段：项目实战（第7-8周）', level=1)
content4 = [
    '综合项目选择（任选其一）：',
    '',
    '项目1：个人博客系统',
    '- 用户注册登录',
    '- 文章发布与管理',
    '- 评论功能',
    '',
    '项目2：数据分析工具',
    '- 数据导入导出',
    '- 数据可视化',
    '- 报表生成',
    '',
    '项目3：自动化办公工具',
    '- Excel处理',
    '- 邮件发送',
    '- 定时任务'
]
for item in content4:
    doc.add_paragraph(item)

# 学习建议
doc.add_heading('学习建议', level=1)
tips = [
    '1. 每天至少练习1小时',
    '2. 遇到问题先自己思考，再查资料',
    '3. 多写代码，多调试',
    '4. 每周总结学习内容',
    '5. 完成每个阶段的练习题目'
]
for tip in tips:
    doc.add_paragraph(tip)

# 保存文档
save_path = r'D:\AI\Python学习计划.docx'
doc.save(save_path)
print(f'Word文档已保存: {save_path}')
