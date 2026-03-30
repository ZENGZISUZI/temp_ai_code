# -*- coding: utf-8 -*-
"""
PDF转Word工具
功能：读取PDF文本内容，生成Word文档
支持：文本提取、图片提取、表格识别
"""

import os
import sys

def check_and_install_packages():
    """检查并安装必要的包"""
    packages = {
        'pdfplumber': 'pdfplumber',
        'python-docx': 'docx',
        'Pillow': 'PIL'
    }
    
    missing = []
    for pip_name, import_name in packages.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pip_name)
    
    if missing:
        print(f"正在安装依赖包: {', '.join(missing)}")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q'] + missing)
        print("安装完成！")

check_and_install_packages()

import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io
from pathlib import Path


def set_chinese_font(doc):
    """设置中文字体"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_heading_with_style(doc, text, level):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def extract_images_from_page(page, page_num, output_dir):
    """从PDF页面提取图片"""
    images = []
    img_dir = output_dir / 'images'
    img_dir.mkdir(exist_ok=True)
    
    try:
        # 获取页面中的图片
        for i, img in enumerate(page.images):
            try:
                # 获取图片数据
                x0, top, x1, bottom = img['x0'], img['top'], img['x1'], img['bottom']
                width = x1 - x0
                height = bottom - top
                
                # 尝试提取图片
                if 'stream' in img:
                    stream = img['stream']
                    if hasattr(stream, 'get_data'):
                        img_data = stream.get_data()
                        if img_data:
                            img_path = img_dir / f'page{page_num}_img{i}.png'
                            with open(img_path, 'wb') as f:
                                f.write(img_data)
                            images.append({
                                'path': str(img_path),
                                'width': width,
                                'height': height
                            })
            except Exception as e:
                continue
    except Exception as e:
        pass
    
    return images


def pdf_to_word(pdf_path, output_path=None, extract_images=True):
    """
    将PDF转换为Word文档
    
    参数:
        pdf_path: PDF文件路径
        output_path: 输出Word路径（可选，默认与PDF同目录）
        extract_images: 是否提取图片
    """
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
    
    if output_path is None:
        output_path = pdf_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    output_dir = output_path.parent
    
    print(f"正在处理: {pdf_path.name}")
    
    # 创建Word文档
    doc = Document()
    set_chinese_font(doc)
    
    # 添加文档标题
    add_heading_with_style(doc, pdf_path.stem, 0)
    doc.add_paragraph()
    
    # 打开PDF
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        print(f"共 {total_pages} 页")
        
        for page_num, page in enumerate(pdf.pages, 1):
            print(f"  处理第 {page_num}/{total_pages} 页...", end='\r')
            
            # 添加页码标记
            if total_pages > 1:
                page_marker = doc.add_paragraph()
                page_marker.add_run(f"—— 第 {page_num} 页 ——").bold = True
                page_marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 提取文本
            text = page.extract_text()
            if text:
                # 按段落分割
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # 判断是否为标题（简单规则）
                        if len(para_text) < 50 and not para_text.endswith(('。', '，', '、', '；', '：')):
                            # 可能是标题
                            p = doc.add_paragraph()
                            run = p.add_run(para_text)
                            run.bold = True
                            run.font.size = Pt(12)
                        else:
                            doc.add_paragraph(para_text)
            
            # 提取表格
            tables = page.extract_tables()
            if tables:
                for table_data in tables:
                    if table_data and len(table_data) > 0:
                        # 创建表格
                        rows = len(table_data)
                        cols = max(len(row) for row in table_data)
                        
                        if rows > 0 and cols > 0:
                            table = doc.add_table(rows=rows, cols=cols)
                            table.style = 'Table Grid'
                            
                            for i, row_data in enumerate(table_data):
                                for j, cell_data in enumerate(row_data):
                                    if j < cols and cell_data:
                                        table.rows[i].cells[j].text = str(cell_data).strip()
                            
                            doc.add_paragraph()  # 表格后加空行
            
            # 提取图片
            if extract_images:
                images = extract_images_from_page(page, page_num, output_dir)
                for img_info in images:
                    try:
                        doc.add_picture(img_info['path'], width=Inches(4))
                        doc.add_paragraph()  # 图片后加空行
                    except Exception as e:
                        pass
            
            # 页间分隔
            if page_num < total_pages:
                doc.add_paragraph()
        
        print(f"\n处理完成！")
    
    # 保存Word文档
    doc.save(output_path)
    print(f"Word文档已保存: {output_path}")
    
    return output_path


def batch_convert(input_dir, output_dir=None):
    """批量转换目录下的所有PDF"""
    input_dir = Path(input_dir)
    
    if not input_dir.exists():
        raise FileNotFoundError(f"目录不存在: {input_dir}")
    
    pdf_files = list(input_dir.glob('*.pdf'))
    
    if not pdf_files:
        print(f"目录中没有找到PDF文件: {input_dir}")
        return
    
    print(f"找到 {len(pdf_files)} 个PDF文件")
    
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
    else:
        output_dir = input_dir
    
    success = 0
    failed = 0
    
    for pdf_file in pdf_files:
        try:
            output_path = output_dir / (pdf_file.stem + '.docx')
            pdf_to_word(pdf_file, output_path)
            success += 1
        except Exception as e:
            print(f"转换失败 {pdf_file.name}: {e}")
            failed += 1
    
    print(f"\n批量转换完成！成功: {success}, 失败: {failed}")


def main():
    """主函数"""
    print("=" * 50)
    print("PDF 转 Word 工具")
    print("=" * 50)
    print()
    
    print("请选择模式:")
    print("1. 转换单个PDF文件")
    print("2. 批量转换目录下所有PDF")
    print()
    
    choice = input("请输入选项 (1/2): ").strip()
    
    if choice == '1':
        pdf_path = input("请输入PDF文件路径: ").strip()
        if not pdf_path:
            print("路径不能为空")
            return
        
        # 处理路径中的引号
        pdf_path = pdf_path.strip('"').strip("'")
        
        output = input("请输入输出Word路径 (直接回车使用默认路径): ").strip().strip('"').strip("'")
        
        try:
            result = pdf_to_word(pdf_path, output if output else None)
            print(f"\n✓ 转换成功！")
            print(f"输出文件: {result}")
        except Exception as e:
            print(f"\n✗ 转换失败: {e}")
    
    elif choice == '2':
        input_dir = input("请输入PDF所在目录: ").strip().strip('"').strip("'")
        output_dir = input("请输入输出目录 (直接回车使用原目录): ").strip().strip('"').strip("'")
        
        try:
            batch_convert(input_dir, output_dir if output_dir else None)
        except Exception as e:
            print(f"\n✗ 批量转换失败: {e}")
    
    else:
        print("无效选项")
    
    input("\n按回车键退出...")


if __name__ == '__main__':
    # ==================== 配置区域 ====================
    # PDF文件路径（修改这里）
    PDF_PATH = r"D:\AI\test.pdf"
    
    # 输出Word路径（None表示与PDF同目录）
    OUTPUT_PATH = None  # 例如: r"D:\AI\output.docx"
    
    # 是否提取图片
    EXTRACT_IMAGES = True
    # ================================================
    
    # 执行转换
    pdf_to_word(PDF_PATH, OUTPUT_PATH, EXTRACT_IMAGES)
